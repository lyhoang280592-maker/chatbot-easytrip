import os
import re
import asyncio
import httpx
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

from generate_contracts import num_to_vietnamese_words, num_to_english_words

load_dotenv("/Users/phamtranthuyvy/Projects/chatbot-easytrip/.env")

LARK_APP_ID = os.getenv("LARK_APP_ID")
LARK_APP_SECRET = os.getenv("LARK_APP_SECRET")
LARK_APP_TOKEN = os.getenv("LARK_APP_TOKEN")
TABLE_ID = "tbluosVi3sQS9gIS"


async def fetch_customer_data_mapping():
    url = "https://open.larksuite.com/open-apis/auth/v3/tenant_access_token/internal"
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.post(url, json={"app_id": LARK_APP_ID, "app_secret": LARK_APP_SECRET})
        token = r.json().get("tenant_access_token")

    headers = {"Authorization": f"Bearer {token}"}
    all_records = []
    page_token = None
    while True:
        p_url = f"https://open.larksuite.com/open-apis/bitable/v1/apps/{LARK_APP_TOKEN}/tables/{TABLE_ID}/records?page_size=500"
        if page_token: p_url += f"&page_token={page_token}"
        async with httpx.AsyncClient(timeout=30) as client:
            r_rec = await client.get(p_url, headers=headers)
            data_res = r_rec.json().get("data", {})
            items = data_res.get("items", [])
            all_records.extend(items)
            if not data_res.get("has_more"): break
            page_token = data_res.get("page_token")

    start_ts = 1785517200000 # 01/08/2026
    end_ts = 1787158799000   # 19/08/2026

    agency_channels = {"sergei", "bolot", "mr.vong", "iuliia sotckaia", "alex"}
    
    mapping = {}
    matched = []
    for it in all_records:
        f = it.get("fields", {})
        ev_ts = f.get("Event Date")
        if ev_ts and start_ts <= ev_ts <= end_ts:
            code_ev = f.get("Code EV")
            ev_file = f.get("EV")
            if code_ev or ev_file:
                matched.append(f)

    # Sắp xếp theo Event Date giảm dần
    matched.sort(key=lambda x: x.get("Event Date") or 0, reverse=True)

    for idx, f in enumerate(matched, 1):
        name = str(f.get("Tên khách (Name)", "") or "").strip().upper()
        ch_list = f.get("Nguồn( Channel)", ["Direct"]) or ["Direct"]
        ch_raw = str(ch_list[0] if ch_list else "Direct")
        ch = ch_raw.strip().lower()
        
        is_agency = ch in agency_channels
        
        srv = str(f.get("Loại dịch vụ (type)", [""])[0] if f.get("Loại dịch vụ (type)") else "")
        note = str(f.get("Ghi chú", "") or "")
        visa_cam = f.get("Visa Cam")
        
        is_cam = ("cambodia" in srv.lower()) or (visa_cam is not None)
        is_multi = ("multi" in note.lower()) or ("multi" in srv.lower())
        
        if is_cam:
            state_fee = 30 * 26500  # 795.000 VNĐ
            usd = 30
            v_type = "CAM"
        elif is_multi:
            state_fee = 50 * 26500  # 1.325.000 VNĐ
            usd = 50
            v_type = "MULTI"
        else:
            state_fee = 25 * 26500  # 662.500 VNĐ
            usd = 25
            v_type = "SINGLE"

        try:
            sales_crm = int(str(f.get("Sales revenue", "0") or "0").replace(".", "").replace(",", ""))
        except:
            sales_crm = 0

        info = {
            "name": name,
            "is_agency": is_agency,
            "channel": ch_raw,
            "state_fee": state_fee,
            "usd": usd,
            "v_type": v_type,
            "sales_crm": sales_crm,
            "stt": idx
        }

        # Lưu mapping theo STT và theo tên
        mapping[f"STT_{idx:02d}"] = info
        if name:
            safe_name = re.sub(r'[^a-zA-Z0-9]', '', name)
            mapping[safe_name] = info

    return mapping


def update_contract_fees(src_path, dst_path, cust_info):
    doc = docx.Document(src_path)
    
    is_agency = cust_info["is_agency"]
    state_fee = cust_info["state_fee"]
    usd_val = cust_info["usd"]
    sales_crm = cust_info["sales_crm"]

    # 1. Tính Tổng cộng và Phí dịch vụ:
    # - Khách lẻ: Tổng cộng = Sales CRM, Phí DV = Sales CRM - Lệ phí Nhà nước
    # - Đại lý: Tổng cộng = Sales CRM * 108%, Phí DV = Tổng cộng - Lệ phí Nhà nước
    if is_agency:
        total_amount = int(round(sales_crm * 1.08)) if sales_crm > 0 else (757500 + state_fee)
    else:
        total_amount = sales_crm if sales_crm > 0 else (757500 + state_fee)

    if total_amount <= state_fee:
        service_fee = 757500
        total_amount = service_fee + state_fee
    else:
        service_fee = total_amount - state_fee

    # 2. Tìm Bảng 2 và cập nhật
    target_table = None
    for tbl in doc.tables:
        header_text = "".join([c.text for c in tbl.rows[0].cells])
        if "Nội dung chi phí" in header_text or "Description of Fees" in header_text:
            target_table = tbl
            break

    if target_table and len(target_table.rows) >= 4:
        # Cập nhật Row 1: Phí dịch vụ Easy Trip
        r1_cells = target_table.rows[1].cells
        for cell_idx in [3, 4]:
            if cell_idx < len(r1_cells):
                r1_cells[cell_idx].text = f"{service_fee:,}".replace(",", ".")
                for p in r1_cells[cell_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(9.0)

        # Cập nhật Row 2: Lệ phí Nhà nước
        r2_cells = target_table.rows[2].cells
        r2_desc_p = r2_cells[1].paragraphs[0]
        r2_desc_p.text = ""
        
        r_vi = r2_desc_p.add_run(f"Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử (Thu hộ - Chi hộ)\n(${usd_val} – tỷ giá 26.500VND)\n")
        r_vi.font.name = "Arial"
        r_vi.font.size = Pt(9.0)
        
        r_en = r2_desc_p.add_run(f"State fee for electronic visa registration (Collected & Paid on behalf)\n(${usd_val} – exchange rate 26,500 VND)")
        r_en.font.name = "Arial"
        r_en.font.size = Pt(8.0)
        r_en.font.italic = True
        r2_desc_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for cell_idx in [3, 4]:
            if cell_idx < len(r2_cells):
                r2_cells[cell_idx].text = f"{state_fee:,}".replace(",", ".")
                for p in r2_cells[cell_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(9.0)

        # Cập nhật Row 3: Tổng cộng
        r3_cells = target_table.rows[3].cells
        if len(r3_cells) > 0:
            r3_cells[-1].text = f"{total_amount:,}".replace(",", ".")
            for p in r3_cells[-1].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    r.font.bold = True

    # 3. Cập nhật các đoạn văn bản ghi số tiền bằng chữ (Bằng chữ / In words)
    vn_words = num_to_vietnamese_words(total_amount)
    en_words = num_to_english_words(total_amount)

    for p in doc.paragraphs:
        if p.text.startswith("Bằng chữ:"):
            txt_vn = vn_words.strip()
            if not txt_vn.endswith("."): txt_vn += "."
            p.text = f"Bằng chữ: {txt_vn}"
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9.5)
                r.font.bold = True
        elif p.text.startswith("In words:"):
            txt_en = en_words.strip()
            if not txt_en.endswith("."): txt_en += "."
            p.text = f"In words: {txt_en}"
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9.0)
                r.font.italic = True

    os.makedirs(os.path.dirname(dst_path), exist_ok=True)
    doc.save(dst_path)
    return total_amount, service_fee, state_fee


def run_all_updates():
    mapping = asyncio.run(fetch_customer_data_mapping())
    src_dir = "HD_Khach_le/HĐ_Khách lẻ_original"
    dst_dir = "HD_Khach_le/HĐ_Khách lẻ"
    
    files = [f for f in sorted(os.listdir(src_dir)) if f.endswith(".docx")]
    print(f"🚀 Bắt đầu cập nhật {len(files)} file theo quy tắc Khách lẻ vs Đại lý (x 108%)...")
    
    agency_count = 0
    retail_count = 0
    
    for filename in files:
        src_path = os.path.join(src_dir, filename)
        dst_path = os.path.join(dst_dir, filename)
        
        stt_match = re.match(r'^(\d+)_Hop_Dong_', filename)
        stt_key = f"STT_{int(stt_match.group(1)):02d}" if stt_match else None
        
        cust_info = mapping.get(stt_key)
        if not cust_info:
            clean_fn = re.sub(r'^\d+_Hop_Dong_', '', filename.replace('.docx', ''))
            clean_fn = re.sub(r'(_chuacochuky|_kocochuky|_chukybikhuat|_cancel_kocochuky)', '', clean_fn)
            clean_key = re.sub(r'[^a-zA-Z0-9]', '', clean_fn.upper())
            for k, v in mapping.items():
                if k in clean_key or clean_key in k:
                    cust_info = v
                    break
                
        if not cust_info:
            cust_info = {
                "name": clean_fn,
                "is_agency": False,
                "channel": "Direct",
                "state_fee": 25 * 26500,
                "usd": 25,
                "v_type": "SINGLE",
                "sales_crm": 1420000
            }
            
        if cust_info["is_agency"]:
            agency_count += 1
        else:
            retail_count += 1
            
        tot, srv, st = update_contract_fees(src_path, dst_path, cust_info)
        
    print(f"\n🎉 HOÀN TẤT CẬP NHẬT 86 HỢP ĐỒNG:")
    print(f"  - Đại lý (x 108%): {agency_count} hợp đồng")
    print(f"  - Khách lẻ (Tổng giữ nguyên): {retail_count} hợp đồng")
    print(f"  - Thư mục lưu trữ: {dst_dir}")


if __name__ == "__main__":
    run_all_updates()
