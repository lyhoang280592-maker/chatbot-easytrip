import os
import re
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Tra cứu thông tin loại visa từ CRM Lark Base
import asyncio, httpx
from dotenv import load_dotenv
load_dotenv("/Users/phamtranthuyvy/Projects/chatbot-easytrip/.env")

LARK_APP_ID = os.getenv("LARK_APP_ID")
LARK_APP_SECRET = os.getenv("LARK_APP_SECRET")
LARK_APP_TOKEN = os.getenv("LARK_APP_TOKEN")
TABLE_ID = "tbluosVi3sQS9gIS"

async def fetch_customer_fee_mapping():
    url = "https://open.larksuite.com/open-apis/auth/v3/tenant_access_token/internal"
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.post(url, json={"app_id": LARK_APP_ID, "app_secret": LARK_APP_SECRET})
        token = r.json().get("tenant_access_token")

    headers = {"Authorization": f"Bearer {token}"}
    all_records = []
    page_token = None
    while True:
        p_url = f"https://open.larksuite.com/open-apis/bitable/v1/apps/{LARK_APP_TOKEN}/tables/{TABLE_ID}/records?page_size=500"
        if page_token: p_url += f"&page_token={page_token}"
        async with httpx.AsyncClient(timeout=30) as client:
            r_rec = await client.get(p_url, headers=headers)
            data_res = r_rec.json().get("data", {})
            items = data_res.get("items", [])
            all_records.extend(items)
            if not data_res.get("has_more"): break
            page_token = data_res.get("page_token")

    mapping = {}
    for it in all_records:
        f = it.get("fields", {})
        name = str(f.get("Tên khách (Name)", "") or "").strip().upper()
        srv = str(f.get("Loại dịch vụ (type)", [""])[0] if f.get("Loại dịch vụ (type)") else "")
        note = str(f.get("Ghi chú", "") or "")
        visa_cam = f.get("Visa Cam")
        
        is_cam = ("cambodia" in srv.lower()) or (visa_cam is not None)
        is_multi = ("multi" in note.lower()) or ("multi" in srv.lower())
        
        if is_cam:
            fee = 30 * 26500  # 795.000 VNĐ
            v_type = "CAM"
            usd = 30
        elif is_multi:
            fee = 50 * 26500  # 1.325.000 VNĐ
            v_type = "MULTI"
            usd = 50
        else:
            fee = 25 * 26500  # 662.500 VNĐ
            v_type = "SINGLE"
            usd = 25

        if name:
            safe_name = re.sub(r'[^a-zA-Z0-9]', '', name)
            mapping[safe_name] = (fee, usd, v_type)

    return mapping


def update_table_2_in_file(src_path, dst_path, fee_info):
    """
    Chỉ sửa duy nhất mục Lệ phí Nhà nước (Row 2) và Phí dịch vụ (Row 1) trong Bảng 2.
    Giữ nguyên 100% toàn bộ phần còn lại của file Word.
    """
    doc = docx.Document(src_path)
    state_fee, usd_val, v_type = fee_info

    # Tìm Bảng 2 (Bảng chi tiết chi phí dịch vụ)
    target_table = None
    for tbl in doc.tables:
        header_text = "".join([c.text for c in tbl.rows[0].cells])
        if "Nội dung chi phí" in header_text or "Description of Fees" in header_text:
            target_table = tbl
            break

    if target_table and len(target_table.rows) >= 4:
        # 1. Đọc tổng tiền hiện tại từ Row 3 (Tổng cộng)
        r3_cells = target_table.rows[3].cells
        total_text = r3_cells[-1].text.strip().replace(".", "").replace(",", "")
        try:
            total_amount = int(total_text)
        except:
            total_amount = 0

        if total_amount <= state_fee:
            service_fee = 757500
            total_amount = service_fee + state_fee
        else:
            service_fee = total_amount - state_fee

        # 2. Cập nhật Row 1: Phí dịch vụ Easy Trip
        r1_cells = target_table.rows[1].cells
        # Unit price & Total for Row 1
        for cell_idx in [3, 4]:
            if cell_idx < len(r1_cells):
                r1_cells[cell_idx].text = f"{service_fee:,}".replace(",", ".")
                for p in r1_cells[cell_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(9.0)

        # 3. Cập nhật Row 2: Lệ phí Nhà nước
        r2_cells = target_table.rows[2].cells
        # Cập nhật nội dung mô tả Row 2
        r2_desc_p = r2_cells[1].paragraphs[0]
        r2_desc_p.text = ""
        
        r_vi = r2_desc_p.add_run(f"Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử (Thu hộ - Chi hộ)\n(${usd_val} – tỷ giá 26.500VND)\n")
        r_vi.font.name = "Arial"
        r_vi.font.size = Pt(9.0)
        
        r_en = r2_desc_p.add_run(f"State fee for electronic visa registration (Collected & Paid on behalf)\n(${usd_val} – exchange rate 26,500 VND)")
        r_en.font.name = "Arial"
        r_en.font.size = Pt(8.0)
        r_en.font.italic = True
        r2_desc_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Unit price & Total for Row 2
        for cell_idx in [3, 4]:
            if cell_idx < len(r2_cells):
                r2_cells[cell_idx].text = f"{state_fee:,}".replace(",", ".")
                for p in r2_cells[cell_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(9.0)

        # 4. Cập nhật Row 3: Tổng cộng (nếu có điều chỉnh)
        if len(r3_cells) > 0:
            r3_cells[-1].text = f"{total_amount:,}".replace(",", ".")
            for p in r3_cells[-1].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    r.font.bold = True

    os.makedirs(os.path.dirname(dst_path), exist_ok=True)
    doc.save(dst_path)


def process_all():
    mapping = asyncio.run(fetch_customer_fee_mapping())
    src_dir = "HD_Khach_le/HĐ_Khách lẻ_original"
    dst_dir = "HD_Khach_le/HĐ_Khách lẻ"
    
    files = [f for f in sorted(os.listdir(src_dir)) if f.endswith(".docx")]
    print(f"🔄 Đang xử lý cập nhật biểu phí cho {len(files)} file gốc từ Google Drive...")
    
    updated_count = 0
    for filename in files:
        src_path = os.path.join(src_dir, filename)
        dst_path = os.path.join(dst_dir, filename)
        
        # Tìm thông tin loại visa từ filename
        clean_fn = re.sub(r'^\d+_Hop_Dong_', '', filename.replace('.docx', ''))
        clean_fn = re.sub(r'(_chuacochuky|_kocochuky|_chukybikhuat|_cancel_kocochuky)', '', clean_fn)
        clean_key = re.sub(r'[^a-zA-Z0-9]', '', clean_fn.upper())
        
        fee_info = (25 * 26500, 25, "SINGLE") # Mặc định
        for k, v in mapping.items():
            if k in clean_key or clean_key in k:
                fee_info = v
                break
                
        update_table_2_in_file(src_path, dst_path, fee_info)
        updated_count += 1
        
    print(f"🎉 Đã hoàn tất cập nhật chính xác {updated_count}/{len(files)} file tại '{dst_dir}'!")


if __name__ == "__main__":
    process_all()
