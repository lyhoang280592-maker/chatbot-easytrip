import os
import re
import json
import asyncio
import httpx
from datetime import datetime
from dotenv import load_dotenv
import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from generate_contracts import (
    num_to_vietnamese_words,
    num_to_english_words,
    extract_and_clean_signature
)

load_dotenv("/Users/phamtranthuyvy/Projects/chatbot-easytrip/.env")

LARK_APP_ID = os.getenv("LARK_APP_ID")
LARK_APP_SECRET = os.getenv("LARK_APP_SECRET")
LARK_APP_TOKEN = os.getenv("LARK_APP_TOKEN")
TABLE_ID = "tbluosVi3sQS9gIS"


def set_cell_border(cell, **kwargs):
    """
    Set cell borders in python-docx
    kwargs: top, bottom, left, right (val='single', sz='4', color='000000')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell margins (padding) in twips (1 pt = 20 twips)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_p(doc, text="", font_name="Arial", font_size=10, bold=False, italic=False, 
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2, line_spacing=1.15):
    """Helper to add styled paragraph with single run or base formatting"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def build_contract_docx(output_docx_path: str, data: dict):
    """
    Tạo tệp Word .docx Hợp đồng 7 trang chuẩn xác 100% theo mẫu
    """
    doc = docx.Document()
    
    # Thiết lập căn lề chuẩn A4 2.0 cm
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)

    # -------------------------------------------------------------
    # TRANG 1: TIÊU NGỮ & BÊN A
    # -------------------------------------------------------------
    add_p(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "Độc lập - Tự do - Hạnh phúc", font_size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "---------------", font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    add_p(doc, "HỢP ĐỒNG CUNG CẤP DỊCH VỤ TƯ VẤN VÀ LÀM HỒ SƠ VISA", font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc, "CONSULTING AND VISA APPLICATION SERVICE CONTRACT", font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, f"Số / No: {data.get('contract_no', '......')}/2026/HĐDV-EASYTRIP", font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    legal_bases = [
        ("Căn cứ Bộ luật Dân sự số 91/2015/QH13 ngày 24/11/2015 và các văn bản hướng dẫn thi hành;",
         "Pursuant to the Civil Code No. 91/2015/QH13 dated November 24, 2015 and its guiding documents;"),
        ("Căn cứ Luật Thương mại số 36/2005/QH11 ngày 14/06/2005 và các văn bản hướng dẫn thi hành;",
         "Pursuant to the Commercial Law No. 36/2005/QH11 dated June 14, 2005 and its guiding documents;"),
        ("Căn cứ Luật Doanh nghiệp số 59/2020/QH14 ngày 17/06/2020 và các văn bản hướng dẫn thi hành;",
         "Pursuant to the Law on Enterprises No. 59/2020/QH14 dated June 17, 2020 and its guiding documents;"),
        ("Căn cứ Luật Nhập cảnh, xuất cảnh, quá cảnh, cư trú của người nước ngoài tại Việt Nam số 47/2014/QH13 ngày 16/06/2014 và các luật sửa đổi, bổ sung liên quan;",
         "Pursuant to the Law on Entry, Exit, Transit, and Residence of Foreigners in Vietnam No. 47/2014/QH13 dated June 16, 2014 and its related amendments;"),
        ("Căn cứ Luật Quản lý thuế số 38/2019/QH14 ngày 13/06/2019 và Luật Thuế giá trị gia tăng số 13/2008/QH12 cùng các văn bản hướng dẫn thi hành;",
         "Pursuant to the Law on Tax Administration No. 38/2019/QH14 dated June 13, 2019 and the Law on Value Added Tax No. 13/2008/QH12 and their guiding documents;"),
        ("Căn cứ Thông tư số 28/2026/TT-BTC ngày 27/03/2026 quy định về phí và lệ phí trong lĩnh vực xuất cảnh, nhập cảnh, quá cảnh, cư trú tại Việt Nam;",
         "Pursuant to Circular No. 28/2026/TT-BTC dated March 27, 2026 regulating fees and charges in the field of entry, exit, transit, and residence in Vietnam;"),
        ("Căn cứ vào nhu cầu và khả năng thực tế của hai Bên.",
         "Pursuant to the demands and actual capacities of both Parties.")
    ]

    for vi, en in legal_bases:
        add_p(doc, f"- {vi}", font_size=9, space_after=1)
        add_p(doc, en, font_size=8.5, italic=True, space_after=4)

    contract_date_vi = data.get('date_vi', 'ngày 16 tháng 06 năm 2026')
    contract_date_en = data.get('date_en', 'June 16, 2026')

    add_p(doc, f"Hôm nay, {contract_date_vi}, tại văn phòng CÔNG TY TNHH CHUYẾN ĐI VÀ THỊ THỰC DỄ DÀNG, chúng tôi gồm có:", font_size=9.5, space_before=4, space_after=1)
    add_p(doc, f"Today, {contract_date_en}, at the office of EASY TRIP AND VISA COMPANY LIMITED, we consist of:", font_size=9, italic=True, space_after=6)

    add_p(doc, "BÊN A: BÊN SỬ DỤNG DỊCH VỤ (KHÁCH HÀNG)", font_size=9.5, bold=True, space_after=1)
    add_p(doc, "PARTY A: SERVICE USER (CUSTOMER)", font_size=9.5, bold=True, space_after=4)

    # Customer info
    p_name = add_p(doc, font_size=9.5, space_after=2)
    p_name.add_run("Tên tổ chức/cá nhân / Name: ")
    r_name = p_name.add_run(str(data.get('customer_name', 'IRNRNAZAROV ENVER')).upper())
    r_name.bold = True

    p_pass = add_p(doc, font_size=9.5, space_after=2)
    p_pass.add_run("Số Hộ chiếu / Passport No.: ")
    r_pass = p_pass.add_run(str(data.get('passport_no', '767587433')))
    r_pass.bold = True
    p_pass.add_run("    Ngày cấp / Date of Issue: ")
    r_doi = p_pass.add_run(str(data.get('date_of_issue', '20/05/2022')))
    r_doi.bold = True

    p_nat = add_p(doc, font_size=9.5, space_after=2)
    p_nat.add_run("Quốc tịch / Nationality: ")
    r_nat = p_nat.add_run(str(data.get('nationality', 'Nga / Russian')))
    r_nat.bold = True

    add_p(doc, f"Email: {data.get('email', '.................................................................................')}", font_size=9.5, space_after=2)
    add_p(doc, "Đại diện bởi / Represented by: .......................................................", font_size=9.5, space_after=4)

    # -------------------------------------------------------------
    # TRANG 2: BÊN B & ĐIỀU 1
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "BÊN B: BÊN CUNG CẤP DỊCH VỤ", font_size=9.5, bold=True, space_after=1)
    add_p(doc, "PARTY B: SERVICE PROVIDER", font_size=9.5, bold=True, space_after=4)

    p_b_name = add_p(doc, font_size=9.5, space_after=1)
    p_b_name.add_run("Tên đơn vị / Company Name: ")
    r_bn = p_b_name.add_run("CÔNG TY TNHH CHUYẾN ĐI VÀ THỊ THỰC DỄ DÀNG")
    r_bn.bold = True

    add_p(doc, "English Name: EASY TRIP AND VISA COMPANY LIMITED", font_size=9, italic=True, space_after=2)
    add_p(doc, "Mã số thuế / Tax Code: 4202051389", font_size=9.5, space_after=2)
    add_p(doc, "Địa chỉ / Address: 21 Phan Vinh, Phường Vĩnh Nguyên, TP. Nha Trang, tỉnh Khánh Hòa, Việt Nam", font_size=9.5, space_after=2)

    p_rep = add_p(doc, font_size=9.5, space_after=2)
    p_rep.add_run("Đại diện / Represented by: Ông/Bà ")
    r_rep = p_rep.add_run("LÝ VIỆT HOÀNG")
    r_rep.bold = True
    p_rep.add_run(" (Mr./Ms. LY VIET HOANG)")

    add_p(doc, "Chức vụ / Position: Giám đốc / Director", font_size=9.5, space_after=2)
    add_p(doc, "Điện thoại / Telephone: 0896916361 - Hotline: 0896916361", font_size=9.5, space_after=2)

    p_acc = add_p(doc, font_size=9.5, space_after=4)
    p_acc.add_run("Tài khoản ngân hàng / Bank Account: VCB Khánh Hòa Số Tài khoản / Account No.: ")
    r_acc = p_acc.add_run("1068582577")
    r_acc.bold = True

    add_p(doc, "Hai Bên cùng thống nhất ký kết Hợp đồng dịch vụ làm thủ tục xin Visa với các điều khoản sau:", font_size=9.5, space_before=4, space_after=1)
    add_p(doc, "Both Parties hereby agree to enter into this Visa Service Contract with the following terms and conditions:", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 1: NỘI DUNG DỊCH VỤ", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 1: SCOPE OF SERVICES", font_size=10, bold=True, space_after=4)

    add_p(doc, "1.1 Bên A giao và Bên B đồng ý thực hiện dịch vụ tư vấn, chuẩn bị, hoàn thiện bộ hồ sơ để nộp hồ sơ và nộp lệ phí xin cấp Visa (Thị thực) cho Bên A với chi tiết sau:", font_size=9.5, space_after=1)
    add_p(doc, "1.1 Party A delegates and Party B agrees to perform the services of consulting preparing, completing, and payment of visa submitting visa applications for Party A as detailed below:", font_size=9, italic=True, space_after=6)

    # Table 1: Scope of Services
    t1 = doc.add_table(rows=2, cols=6)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths1 = [Cm(1.0), Cm(3.8), Cm(2.5), Cm(4.2), Cm(2.2), Cm(2.8)]

    headers1 = [
        ("STT\nNo.", True),
        ("Họ và tên người xin Visa\nFull name of applicant", True),
        ("Số Hộ chiếu\nPassport No.", True),
        ("Cơ quan cấp thị thực\nIssuing Authority", True),
        ("Mục đích\nPurpose", True),
        ("Loại Visa\nVisa Type", True)
    ]

    for c_idx, (h_text, is_bold) in enumerate(headers1):
        cell = t1.cell(0, c_idx)
        cell.width = col_widths1[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = h_text.split("\n")
        r1 = p.add_run(lines[0])
        r1.font.name = "Arial"
        r1.font.size = Pt(8.5)
        r1.font.bold = True
        if len(lines) > 1:
            p.add_run("\n")
            r2 = p.add_run(lines[1])
            r2.font.name = "Arial"
            r2.font.size = Pt(8.0)
            r2.font.italic = True

    row1_data = [
        "1",
        str(data.get('customer_name', 'IRNRNAZAROV ENVER')).upper(),
        str(data.get('passport_no', '767587433')),
        "Cục QL XNK - Bộ Công an\nImmigration Dept - MPS (Vietnam)",
        "Du lịch\nTourism",
        "Thị thực điện tử /\nEVisa"
    ]

    for c_idx, val in enumerate(row1_data):
        cell = t1.cell(1, c_idx)
        cell.width = col_widths1[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = val.split("\n")
        r1 = p.add_run(lines[0])
        r1.font.name = "Arial"
        r1.font.size = Pt(8.5)
        if c_idx == 1:
            r1.font.bold = True
        if len(lines) > 1:
            p.add_run("\n")
            r2 = p.add_run(lines[1])
            r2.font.name = "Arial"
            r2.font.size = Pt(8.0)
            r2.font.italic = True

    add_p(doc, space_after=6)
    add_p(doc, "1.2 Nội dung công việc chi tiết của Bên B:", font_size=9.5, space_after=1)
    add_p(doc, "1.2 Detailed scope of work of Party B:", font_size=9, italic=True, space_after=4)
    add_p(doc, "- Tư vấn các quy định pháp lý, điều kiện và thủ tục xin cấp thị thực điện tử của Cục Quản lý Xuất nhập cảnh - Bộ Công an.", font_size=9.5, space_after=1)
    add_p(doc, "Consult on legal regulations, conditions, and procedures for e-visa issuance of the Immigration Department - Ministry of Public Security.", font_size=9, italic=True, space_after=4)
    add_p(doc, "- Hướng dẫn Bên A chuẩn bị giấy tờ, thông tin cá nhân đúng tiêu chuẩn theo yêu cầu.", font_size=9.5, space_after=1)

    # -------------------------------------------------------------
    # TRANG 3: ĐIỀU 2 (CHI PHÍ & THANH TOÁN)
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "Guide Party A to prepare documents and personal information up to standard as required.", font_size=9, italic=True, space_after=4)
    add_p(doc, "- Thực hiện khai trực tuyến thông tin, điền tờ khai và nộp hồ sơ đăng ký.", font_size=9.5, space_after=1)
    add_p(doc, "Perform online declaration, fill out application forms, and submit registration documents.", font_size=9, italic=True, space_after=4)
    add_p(doc, "- Theo dõi tiến trình xử lý của Bộ Công an và bàn giao kết quả thị thực điện tử cho Bên A ngay sau khi có kết quả.", font_size=9.5, space_after=1)
    add_p(doc, "Track the processing status by the Ministry of Public Security and hand over the electronic visa results to Party A immediately upon availability.", font_size=9, italic=True, space_after=8)

    add_p(doc, "ĐIỀU 2: GIÁ TRỊ HỢP ĐỒNG VÀ PHƯƠNG THỨC THANH TOÁN", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 2: CONTRACT VALUE AND PAYMENT METHOD", font_size=10, bold=True, space_after=4)
    add_p(doc, "2.1 Bảng chi tiết chi phí / Detailed Cost Breakdown:", font_size=9.5, space_after=6)

    service_fee = int(data.get('service_fee', 757500))
    state_fee = int(data.get('state_fee', 662500))
    total_amount = service_fee + state_fee

    # Table 2: Cost breakdown
    t2 = doc.add_table(rows=4, cols=5)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths2 = [Cm(1.0), Cm(7.5), Cm(1.6), Cm(3.2), Cm(3.2)]

    headers2 = [
        "STT\nNo.",
        "Nội dung chi phí\nDescription of Fees",
        "Số lượng\nQty",
        "Đơn giá (VNĐ)\nUnit Price (VND)",
        "Thành tiền (VNĐ)\nTotal Amount (VND)"
    ]

    for c_idx, h_text in enumerate(headers2):
        cell = t2.cell(0, c_idx)
        cell.width = col_widths2[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = h_text.split("\n")
        r1 = p.add_run(lines[0])
        r1.font.name = "Arial"
        r1.font.size = Pt(8.5)
        r1.font.bold = True
        if len(lines) > 1:
            p.add_run("\n")
            r2 = p.add_run(lines[1])
            r2.font.name = "Arial"
            r2.font.size = Pt(8.0)
            r2.font.italic = True

    # Row 1: Phí dịch vụ
    r1_items = [
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        ("Phí dịch vụ tư vấn & làm hồ sơ Visa\nService fee for consulting & visa processing", WD_ALIGN_PARAGRAPH.LEFT),
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        (f"{service_fee:,}".replace(",", "."), WD_ALIGN_PARAGRAPH.CENTER),
        (f"{service_fee:,}".replace(",", "."), WD_ALIGN_PARAGRAPH.CENTER)
    ]
    for c_idx, (text_val, align_val) in enumerate(r1_items):
        cell = t2.cell(1, c_idx)
        cell.width = col_widths2[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = align_val
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = text_val.split("\n")
        r = p.add_run(lines[0])
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        if len(lines) > 1:
            p.add_run("\n")
            r_sub = p.add_run(lines[1])
            r_sub.font.name = "Arial"
            r_sub.font.size = Pt(8.0)
            r_sub.font.italic = True

    # Row 2: Lệ phí nhà nước
    usd_val = int(round(state_fee / 26500))
    if usd_val <= 0: usd_val = 25
    r2_desc = f"Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử (Thu hộ - Chi hộ)\n(${usd_val} – tỷ giá 26.500VND)\nState fee for electronic visa registration (Collected & Paid on behalf)\n(${usd_val} – exchange rate 26,500 VND)"

    r2_items = [
        ("2", WD_ALIGN_PARAGRAPH.CENTER),
        (r2_desc, WD_ALIGN_PARAGRAPH.LEFT),
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        (f"{state_fee:,}".replace(",", "."), WD_ALIGN_PARAGRAPH.CENTER),
        (f"{state_fee:,}".replace(",", "."), WD_ALIGN_PARAGRAPH.CENTER)
    ]
    for c_idx, (text_val, align_val) in enumerate(r2_items):
        cell = t2.cell(2, c_idx)
        cell.width = col_widths2[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = align_val
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = text_val.split("\n")
        for l_idx, line in enumerate(lines):
            if l_idx > 0: p.add_run("\n")
            r = p.add_run(line)
            r.font.name = "Arial"
            r.font.size = Pt(8.5 if l_idx < 2 else 8.0)
            if l_idx >= 2: r.font.italic = True

    # Row 3: Merge 0..3 and Total
    cell_tot_lbl = t2.cell(3, 0)
    cell_tot_lbl.merge(t2.cell(3, 3))
    set_cell_border(cell_tot_lbl)
    set_cell_margins(cell_tot_lbl, 120, 120, 150, 150)
    p_tot_lbl = cell_tot_lbl.paragraphs[0]
    p_tot_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tot_lbl.paragraph_format.space_before = Pt(0)
    p_tot_lbl.paragraph_format.space_after = Pt(0)
    r_tl = p_tot_lbl.add_run("Tổng cộng / Total Amount")
    r_tl.font.name = "Arial"
    r_tl.font.size = Pt(8.5)
    r_tl.font.bold = True

    cell_tot_val = t2.cell(3, 4)
    cell_tot_val.width = col_widths2[4]
    set_cell_border(cell_tot_val)
    set_cell_margins(cell_tot_val, 120, 120, 150, 150)
    p_tot_val = cell_tot_val.paragraphs[0]
    p_tot_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tot_val.paragraph_format.space_before = Pt(0)
    p_tot_val.paragraph_format.space_after = Pt(0)
    r_tv = p_tot_val.add_run(f"{total_amount:,}".replace(",", "."))
    r_tv.font.name = "Arial"
    r_tv.font.size = Pt(8.5)
    r_tv.font.bold = True

    add_p(doc, space_after=4)

    words_vi = num_to_vietnamese_words(total_amount)
    words_en = num_to_english_words(total_amount)

    p_wvi = add_p(doc, font_size=9.5, space_after=1)
    p_wvi.add_run("Bằng chữ: ")
    r_wvi = p_wvi.add_run(words_vi)
    r_wvi.bold = True

    p_wen = add_p(doc, font_size=9, italic=True, space_after=6)
    p_wen.add_run("In words: ")
    p_wen.add_run(words_en)

    add_p(doc, "2.2 Thuế Giá trị gia tăng (VAT) / Value Added Tax (VAT):", font_size=9.5, bold=True, space_after=2)
    add_p(doc, "- Đơn giá phí dịch vụ nêu trên đã bao gồm thuế GTGT theo quy định của pháp luật Việt Nam. Bên B có trách nhiệm lập và xuất hóa đơn điện tử giá trị gia tăng (GTGT) hợp pháp cho Bên A đối với phần Phí dịch vụ tư vấn & làm hồ sơ trong vòng 03 ngày làm việc kể từ ngày hoàn thành dịch vụ hoặc khi Bên A thanh toán đầy đủ, tùy điều kiện nào đến trước. Bên A có trách nhiệm cung cấp đầy đủ và chính xác thông tin xuất hóa đơn.", font_size=9.5, space_after=2)
    add_p(doc, "- The service fee specified above is inclusive of VAT in accordance with Vietnamese tax laws. Party B is responsible for issuing a valid electronic Value Added Tax (VAT) invoice to Party A for the service fee component within 03 working days from the completion of the service or upon full payment by Party A, whichever comes first. Party A is responsible for providing complete and accurate billing information.", font_size=9, italic=True, space_after=4)

    # -------------------------------------------------------------
    # TRANG 4: ĐIỀU 3 & ĐIỀU 4
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "- Đối với phần Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử (Thu hộ - Chi hộ), Bên B sẽ cung cấp chứng từ, biên lai lệ phí của Bộ Công an/Cổng dịch vụ công cho Bên A và phần này không chịu thuế GTGT của Bên B.", font_size=9.5, space_after=2)
    add_p(doc, "- For the State fee for electronic visa registration (Collected & Paid on behalf), Party B shall provide the official receipt from the Ministry of Public Security / Public Service Portal to Party A, and this component is not subject to VAT by Party B.", font_size=9, italic=True, space_after=4)

    add_p(doc, "2.3 Tiến độ thanh toán / Payment Schedule:", font_size=9.5, bold=True, space_after=1)
    add_p(doc, "- Bên A thanh toán 100% tổng giá trị hợp đồng ngay sau khi ký kết Hợp đồng này.", font_size=9.5, space_after=1)
    add_p(doc, "- Party A shall pay 100% of the total contract value immediately upon signing this Contract.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 3: QUYỀN VÀ NGHĨA VỤ CỦA BÊN A", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 3: RIGHTS AND OBLIGATIONS OF PARTY A", font_size=10, bold=True, space_after=4)

    add_p(doc, "3.1 Cung cấp đầy đủ, trung thực, chính xác và kịp thời các hồ sơ, thông tin, giấy tờ theo hướng dẫn của Bên B. Bên A tự chịu hoàn toàn trách nhiệm pháp lý trước pháp luật về tính chân thật, hợp pháp của các tài liệu đã cung cấp.", font_size=9.5, space_after=1)
    add_p(doc, "3.1 Provide fully, honestly, accurately, and timely all dossiers, information, and documents as guided by Party B. Party A shall be solely and fully liable under the law for the authenticity and legality of the provided documents.", font_size=9, italic=True, space_after=4)

    add_p(doc, "3.2 Thanh toán đầy đủ và đúng hạn các khoản phí theo quy định tại Điều 2 của Hợp đồng này.", font_size=9.5, space_after=1)
    add_p(doc, "3.2 Pay fully and on time all fees specified in Article 2 of this Contract.", font_size=9, italic=True, space_after=4)

    add_p(doc, "3.3 Có mặt đúng giờ hoặc phối hợp cung cấp sinh trắc học, thông tin bổ sung khi có yêu cầu từ Cục Quản lý Xuất nhập cảnh - Bộ Công an hoặc cơ quan có thẩm quyền của Nhà nước.", font_size=9.5, space_after=1)
    add_p(doc, "3.3 Be present on time or coordinate to provide biometrics and additional information upon request from the Immigration Department - Ministry of Public Security or competent State authorities.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 4: QUYỀN VÀ NGHĨA VỤ CỦA BÊN B", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 4: RIGHTS AND OBLIGATIONS OF PARTY B", font_size=10, bold=True, space_after=4)

    add_p(doc, "4.1 Được quyền sử dụng thông tin trên hộ chiếu, thông tin khác của bên A để thực hiện công việc tư vấn và chuẩn bị hồ sơ đăng ký các dịch vụ công của Nhà Nước Việt Nam.", font_size=9.5, space_after=1)
    add_p(doc, "4.1 To be authorized to use Party A's passport information and other information to perform consultancy services and prepare application dossiers for public services of the State of Vietnam.", font_size=9, italic=True, space_after=4)

    add_p(doc, "4.2 Bảo quản cẩn thận, an toàn các giấy tờ gốc, thông tin cá nhân do Bên A bàn giao trong quá trình thực hiện hợp đồng.", font_size=9.5, space_after=1)
    add_p(doc, "4.2 Carefully and safely preserve the original documents and personal information handed over by Party A during the contract execution.", font_size=9, italic=True, space_after=4)

    add_p(doc, "4.3 Bảo mật tuyệt đối mọi thông tin cá nhân, thông tin hồ sơ của Bên A và không tiết lộ cho bất kỳ bên thứ ba nào khi chưa có sự đồng ý bằng văn bản của Bên A, trừ trường hợp cung cấp cho Cục Quản lý Xuất nhập cảnh - Bộ Công an để đăng ký cấp thị thực điện tử hoặc theo yêu cầu của pháp luật.", font_size=9.5, space_after=1)
    add_p(doc, "4.3 Keep strictly confidential all personal and dossier information of Party A and not disclose it to any third party without Party A's prior written consent, except", font_size=9, italic=True, space_after=4)

    # -------------------------------------------------------------
    # TRANG 5: ĐIỀU 5 & ĐIỀU 6
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "for submission to the Immigration Department - Ministry of Public Security for e-visa registration or as required by law.", font_size=9, italic=True, space_after=4)
    add_p(doc, "4.4 Thông báo kịp thời cho Bên A về tiến độ và kết quả hồ sơ. Bàn giao đầy đủ kết quả dịch vụ công ngay sau khi được Nhà Nước Việt Nam ban hành.", font_size=9.5, space_after=1)
    add_p(doc, "4.4 Promptly notify Party A of the progress and results of the application. Fully hand over the results of the public service immediately upon issuance by the State of Vietnam.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 5: ĐIỀU KHOẢN VỀ KẾT QUẢ DỊCH VỤ VÀ HOÀN PHÍ", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 5: SERVICE RESULTS AND REFUND POLICY", font_size=10, bold=True, space_after=4)

    add_p(doc, "5.1 Hai Bên hiểu rõ rằng quyền quyết định cấp hoặc từ chối cấp thị thực điện tử hoàn toàn thuộc thẩm quyền của Cục Quản lý Xuất nhập cảnh - Bộ Công an nước Cộng hòa Xã hội Chủ nghĩa Việt Nam. Bên B không quyết định và không bảo đảm tuyệt đối kết quả cấp thị thực.", font_size=9.5, space_after=1)
    add_p(doc, "5.1 Both Parties clearly understand that the decision to grant or refuse the electronic visa belongs solely to the authority of the Immigration Department - Ministry of Public Security of the Socialist Republic of Vietnam. Party B does not decide and does not guarantee absolute visa issuance results.", font_size=9, italic=True, space_after=4)

    add_p(doc, "5.2 Xử lý khi rớt thị thực do lỗi khách quan (từ phía Cục Quản lý Xuất nhập cảnh - Bộ Công an từ chối mà không do lỗi của Bên nào):", font_size=9.5, space_after=1)
    add_p(doc, "5.2 Handling visa rejection due to objective reasons (rejection by the Immigration Department - Ministry of Public Security without fault of either Party):", font_size=9, italic=True, space_after=4)

    add_p(doc, "- Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử sẽ không được hoàn lại (theo quy định của Bộ Công an và Thông tư số 28/2026/TT-BTC).", font_size=9.5, space_after=1)
    add_p(doc, "The State fee for electronic visa registration will not be refunded (in accordance with the Ministry of Public Security regulations and Circular No. 28/2026/TT-BTC).", font_size=9, italic=True, space_after=4)

    add_p(doc, f"- Bên B sẽ hoàn trả lại 100% Phí dịch vụ tư vấn & làm hồ sơ ({service_fee:,} VNĐ) cho Bên A trong vòng 05 ngày làm việc kể từ ngày nhận được thông báo từ chối cấp thị thực.".replace(",", "."), font_size=9.5, space_after=1)
    add_p(doc, f"Party B shall refund 100% of the Service Fee (VND {service_fee:,}) to Party A within 05 working days from the date of receiving the visa rejection notice.".replace(",", "."), font_size=9, italic=True, space_after=4)

    add_p(doc, "5.3 Trường hợp Bên A đơn phương hủy hợp đồng sau khi Bên B đã tiến hành xử lý hồ sơ hoặc nộp lệ phí, phí dịch vụ và Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử sẽ không được hoàn lại.", font_size=9.5, space_after=1)
    add_p(doc, "5.3 In case Party A unilaterally terminates the contract after Party B has commenced processing the dossier or paid the fees, the service fee and the State fee for electronic visa registration shall not be refunded.", font_size=9, italic=True, space_after=4)

    add_p(doc, "5.4 Trường hợp Bên A cung cấp thông tin, tài liệu giả mạo hoặc sai sự thật dẫn đến việc hồ sơ bị từ chối hoặc bị xử lý theo pháp luật:", font_size=9.5, space_after=1)
    add_p(doc, "5.4 In case Party A provides forged or untruthful information or documents, leading to visa rejection or legal actions:", font_size=9, italic=True, space_after=4)

    add_p(doc, "- Bên B có quyền đơn phương chấm dứt hợp đồng ngay lập tức.", font_size=9.5, space_after=1)
    add_p(doc, "Party B has the right to unilaterally terminate the contract immediately.", font_size=9, italic=True, space_after=4)

    add_p(doc, "- Bên A không được hoàn lại bất kỳ khoản phí nào và phải tự chịu hoàn toàn trách nhiệm trước pháp luật.", font_size=9.5, space_after=1)
    add_p(doc, "Party A shall not be refunded any fees and must bear full legal responsibility.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 6: BẤT KHẢ KHÁNG VÀ GIỚI HẠN TRÁCH NHIỆM", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 6: FORCE MAJEURE AND LIMITATION OF LIABILITY", font_size=10, bold=True, space_after=4)

    # -------------------------------------------------------------
    # TRANG 6: ĐIỀU 7 & ĐIỀU 8
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "6.1 Sự kiện bất khả kháng là sự kiện xảy ra một cách khách quan, không thể lường trước được và không thể khắc phục được mặc dù đã áp dụng mọi biện pháp cần thiết và khả năng cho phép, bao gồm nhưng không giới hạn ở: thiên tai, dịch bệnh, chiến tranh, sự thay đổi đột ngột về chính sách xuất nhập cảnh của Chính phủ Việt Nam hoặc quốc gia liên quan, việc Cục Quản lý Xuất nhập cảnh - Bộ Công an tạm ngừng hoạt động hoặc đóng cửa biên giới.", font_size=9.5, space_after=1)
    add_p(doc, "6.1 A Force Majeure event is an event that occurs objectively, unpredictably, and irremediably despite all necessary and permissible measures being taken, including but not limited to: natural disasters, epidemics, wars, sudden changes in immigration policies of the Vietnamese Government or related countries, temporary suspension of operations of the Immigration Department - Ministry of Public Security, or border closures.", font_size=9, italic=True, space_after=4)

    add_p(doc, "6.2 Trong trường hợp xảy ra Sự kiện bất khả kháng dẫn đến việc chậm trễ hoặc không thể thực hiện nghĩa vụ hợp đồng, bên bị ảnh hưởng sẽ được miễn trừ trách nhiệm và không phải bồi thường thiệt hại, với điều kiện phải thông báo cho bên kia bằng văn bản trong vòng 03 ngày kể từ ngày xảy ra sự kiện. Hai bên sẽ cùng thương lượng để tìm giải pháp khắc phục.", font_size=9.5, space_after=1)
    add_p(doc, "6.2 In the event of a Force Majeure event leading to delay or inability to perform contractual obligations, the affected party shall be exempted from liability and shall not compensate for damages, provided that the other party is notified in writing within 03 days from the occurrence of the event. Both parties shall negotiate to find a remedy.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 7: GIẢI QUYẾT TRANH CHẤP", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 7: DISPUTE RESOLUTION", font_size=10, bold=True, space_after=4)

    add_p(doc, "7.1 Mọi tranh chấp phát sinh từ hoặc liên quan đến Hợp đồng này trước hết sẽ được hai Bên giải quyết thông qua thương lượng, hòa giải trên tinh thần hợp tác và cùng có lợi.", font_size=9.5, space_after=1)
    add_p(doc, "7.1 Any dispute arising from or related to this Contract shall first be resolved by both Parties through negotiation and conciliation in a cooperative and mutually beneficial spirit.", font_size=9, italic=True, space_after=4)

    add_p(doc, "7.2 Trường hợp tranh chấp không thể giải quyết bằng thương lượng trong vòng 30 ngày kể từ ngày phát sinh, một trong hai Bên có quyền đưa vụ việc ra giải quyết tại Tòa án nhân dân có thẩm quyền tại Thành phố Nha Trang, tỉnh Khánh Hòa, Việt Nam để giải quyết theo quy định của pháp luật Việt Nam.", font_size=9.5, space_after=1)
    add_p(doc, "7.2 In case the dispute cannot be resolved through negotiation within 30 days from its occurrence, either Party has the right to refer the dispute to the competent People's Court in Nha Trang City, Khanh Hoa Province, Vietnam for resolution in accordance with Vietnamese law.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 8: ĐIỀU KHOẢN CHUNG", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 8: GENERAL PROVISIONS", font_size=10, bold=True, space_after=4)

    add_p(doc, "8.1 Hợp đồng này được điều chỉnh và giải thích theo quy định của pháp luật nước Cộng hòa Xã hội Chủ nghĩa Việt Nam.", font_size=9.5, space_after=1)
    add_p(doc, "8.1 This Contract shall be governed by and construed in accordance with the laws of the Socialist Republic of Vietnam.", font_size=9, italic=True, space_after=4)

    add_p(doc, "8.2 Hợp đồng này có hiệu lực kể từ ngày ký và tự động thanh lý sau khi Bên B bàn giao kết quả thị thực điện tử cho Bên A và Bên A hoàn tất nghĩa vụ thanh toán.", font_size=9.5, space_after=1)
    add_p(doc, "8.2 This Contract shall take effect from the date of signing and shall be automatically liquidated after Party B hands over the electronic visa results to Party A and Party A completes the payment obligations.", font_size=9, italic=True, space_after=6)

    # -------------------------------------------------------------
    # TRANG 7: KÝ TÊN & ĐÓNG DẤU
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "8.3 Hợp đồng được lập thành 02 (hai) bản bằng tiếng Việt và tiếng Anh. Mỗi bên giữ 01 (một) bản có giá trị pháp lý như nhau. Bản tiếng Việt sẽ được ưu tiên áp dụng nếu có bất kỳ sự khác biệt nào về cách giải thích giữa hai ngôn ngữ.", font_size=9.5, space_after=1)
    add_p(doc, "8.3 The Contract is prepared in 02 (two) copies in both Vietnamese and English. Each party shall keep 01 (one) copy with equal legal validity. The Vietnamese version shall prevail in case of any discrepancy in interpretation between the two languages.", font_size=9, italic=True, space_after=20)

    # Bảng chữ ký 2 cột
    t_sig = doc.add_table(rows=3, cols=2)
    t_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths_sig = [Cm(8.5), Cm(8.5)]

    # Header row
    c_a_hdr = t_sig.cell(0, 0)
    c_a_hdr.width = col_widths_sig[0]
    p_ah = c_a_hdr.paragraphs[0]
    p_ah.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ah1 = p_ah.add_run("ĐẠI DIỆN BÊN A\n")
    r_ah1.bold = True
    r_ah1.font.name = "Arial"
    r_ah1.font.size = Pt(9.5)
    r_ah2 = p_ah.add_run("REPRESENTATIVE OF PARTY A")
    r_ah2.italic = True
    r_ah2.font.name = "Arial"
    r_ah2.font.size = Pt(8.5)

    c_b_hdr = t_sig.cell(0, 1)
    c_b_hdr.width = col_widths_sig[1]
    p_bh = c_b_hdr.paragraphs[0]
    p_bh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bh1 = p_bh.add_run("ĐẠI DIỆN BÊN B\n")
    r_bh1.bold = True
    r_bh1.font.name = "Arial"
    r_bh1.font.size = Pt(9.5)
    r_bh2 = p_bh.add_run("CÔNG TY TNHH CHUYẾN ĐI VÀ THỊ THỰC DỄ DÀNG\n")
    r_bh2.bold = True
    r_bh2.font.name = "Arial"
    r_bh2.font.size = Pt(9.5)
    r_bh3 = p_bh.add_run("REPRESENTATIVE OF PARTY B\nEASY TRIP AND VISA CO., LTD")
    r_bh3.italic = True
    r_bh3.font.name = "Arial"
    r_bh3.font.size = Pt(8.5)

    # Signature row
    c_a_sig = t_sig.cell(1, 0)
    c_a_sig.width = col_widths_sig[0]
    p_as = c_a_sig.paragraphs[0]
    p_as.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_as.paragraph_format.space_before = Pt(8)
    p_as.paragraph_format.space_after = Pt(8)

    sig_path = data.get('signature_image_path')
    if sig_path and os.path.exists(sig_path):
        try:
            p_as.add_run().add_picture(sig_path, width=Cm(3.8))
        except Exception:
            p_as.add_run("\n\n\n")
    else:
        p_as.add_run("\n\n\n")

    c_b_sig = t_sig.cell(1, 1)
    c_b_sig.width = col_widths_sig[1]
    p_bs = c_b_sig.paragraphs[0]
    p_bs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bs.paragraph_format.space_before = Pt(8)
    p_bs.paragraph_format.space_after = Pt(8)
    p_bs.add_run("\n\n\n")

    # Name row
    c_a_name = t_sig.cell(2, 0)
    c_a_name.width = col_widths_sig[0]
    p_an = c_a_name.paragraphs[0]
    p_an.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ann = p_an.add_run(str(data.get('customer_name', 'IRNRNAZAROV ENVER')).upper())
    r_ann.bold = True
    r_ann.font.name = "Arial"
    r_ann.font.size = Pt(9.5)

    c_b_name = t_sig.cell(2, 1)
    c_b_name.width = col_widths_sig[1]
    p_bn = c_b_name.paragraphs[0]
    p_bn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bnn = p_bn.add_run("LÝ VIỆT HOÀNG")
    r_bnn.bold = True
    r_bnn.font.name = "Arial"
    r_bnn.font.size = Pt(9.5)

    os.makedirs(os.path.dirname(output_docx_path) or ".", exist_ok=True)
    doc.save(output_docx_path)
    return output_docx_path


async def generate_all_86_docx(out_dir="output_docx_contracts"):
    """Tải toàn bộ 86 khách hàng có Evisa và tạo 86 tệp .docx chuẩn"""
    url = "https://open.larksuite.com/open-apis/auth/v3/tenant_access_token/internal"
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.post(url, json={"app_id": LARK_APP_ID, "app_secret": LARK_APP_SECRET})
        token = r.json().get("tenant_access_token")

    if not token:
        print("❌ Không lấy được Lark Token!")
        return

    headers = {"Authorization": f"Bearer {token}"}
    all_records = []
    page_token = None
    while True:
        p_url = f"https://open.larksuite.com/open-apis/bitable/v1/apps/{LARK_APP_TOKEN}/tables/{TABLE_ID}/records?page_size=500"
        if page_token: p_url += f"&page_token={page_token}"
        async with httpx.AsyncClient(timeout=30) as client:
            r_rec = await client.get(p_url, headers=headers)
            data_res = r_rec.json().get("data", {})
            items = data_res.get("items", [])
            all_records.extend(items)
            if not data_res.get("has_more"): break
            page_token = data_res.get("page_token")

    start_ts = datetime(2026, 8, 1, 0, 0, 0).timestamp() * 1000
    end_ts = datetime(2026, 8, 19, 23, 59, 59).timestamp() * 1000

    matched = []
    for it in all_records:
        f = it.get("fields", {})
        ev_ts = f.get("Event Date")
        if ev_ts and start_ts <= ev_ts <= end_ts:
            code_ev = f.get("Code EV")
            ev_file = f.get("EV")
            if code_ev or ev_file:
                matched.append(f)

    # Sắp xếp theo Event Date giảm dần
    matched.sort(key=lambda x: x.get("Event Date") or 0, reverse=True)

    os.makedirs(out_dir, exist_ok=True)
    os.makedirs("downloads/passports", exist_ok=True)
    os.makedirs("extracted_signatures", exist_ok=True)

    print(f"🚀 Bắt đầu tạo 86 hợp đồng .docx trong thư mục '{out_dir}'...")

    generated_files = []

    async with httpx.AsyncClient(timeout=30) as client:
        for idx, c in enumerate(matched, 1):
            name = str(c.get("Tên khách (Name)", "CUSTOMER") or "CUSTOMER").strip()
            nat = c.get("Quốc tịch (National)", ["Nga / Russian"])[0] if c.get("Quốc tịch (National)") else "Nga / Russian"
            code_ev = str(c.get("Code EV", "") or "").strip()
            
            # Tính lệ phí Nhà nước theo loại visa:
            # - Visa Single: $25 * 26.500 = 662.500 VNĐ
            # - Visa Multi:  $50 * 26.500 = 1.325.000 VNĐ
            # - Visa Cam:    $30 * 26.500 = 795.000 VNĐ
            srv_type = str(c.get("Loại dịch vụ (type)", [""])[0] if c.get("Loại dịch vụ (type)") else "")
            note_str = str(c.get("Ghi chú", "") or "")
            visa_cam_val = c.get("Visa Cam")
            
            is_cam = ("cambodia" in srv_type.lower()) or (visa_cam_val is not None)
            is_multi = ("multi" in note_str.lower()) or ("multi" in srv_type.lower())
            
            if is_cam:
                state_fee = 30 * 26500   # 795.000 VNĐ (Visa Cam)
            elif is_multi:
                state_fee = 50 * 26500   # 1.325.000 VNĐ (Visa Multi)
            else:
                state_fee = 25 * 26500   # 662.500 VNĐ (Visa Single)

            try:
                total_rev = int(str(c.get("Sales revenue", "0") or "0").replace(".", "").replace(",", ""))
            except:
                total_rev = 0

            if total_rev <= state_fee:
                service_fee = 757500
                total_rev = service_fee + state_fee
            else:
                service_fee = total_rev - state_fee

            ev_ts = c.get("Event Date") or c.get("Apply Date")
            if ev_ts:
                dt = datetime.fromtimestamp(ev_ts / 1000)
                date_vi = dt.strftime("ngày %d tháng %m năm %Y")
                date_en = dt.strftime("%B %d, %Y")
            else:
                date_vi = "ngày 16 tháng 06 năm 2026"
                date_en = "June 16, 2026"

            safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', name.strip())
            passport_match = re.search(r'[A-Za-z0-9]{7,9}$', code_ev)
            passport_no = passport_match.group(0) if passport_match else (code_ev.replace("E26", "")[:9] if code_ev else "767587433")

            # Xử lý tải ảnh hộ chiếu và trích xuất chữ ký (nếu có)
            sig_path = os.path.join("extracted_signatures", f"{safe_name}_{passport_no}_sig.png")
            passports = c.get("Ảnh hộ chiếu")
            if not os.path.exists(sig_path) and passports and isinstance(passports, list) and len(passports) > 0:
                p_item = passports[0]
                ft = p_item.get("file_token")
                p_path = os.path.join("downloads/passports", f"{safe_name}_{passport_no}.jpg")
                if not os.path.exists(p_path) and ft:
                    try:
                        dl_url = f"https://open.larksuite.com/open-apis/drive/v1/medias/{ft}/download"
                        r_dl = await client.get(dl_url, headers=headers)
                        if r_dl.status_code == 200:
                            with open(p_path, "wb") as img_f:
                                img_f.write(r_dl.content)
                    except Exception:
                        pass
                if os.path.exists(p_path):
                    extract_and_clean_signature(p_path, sig_path)

            contract_no = str(abs(hash(name)) % 900 + 100).zfill(3)

            contract_data = {
                "contract_no": contract_no,
                "date_vi": date_vi,
                "date_en": date_en,
                "customer_name": name.upper(),
                "passport_no": passport_no,
                "date_of_issue": "20/05/2022",
                "nationality": nat,
                "service_fee": service_fee,
                "state_fee": state_fee,
                "signature_image_path": sig_path if os.path.exists(sig_path) else None
            }

            filename = f"{idx:02d}_Hop_Dong_{safe_name}.docx"
            file_path = os.path.join(out_dir, filename)
            build_contract_docx(file_path, contract_data)
            generated_files.append(file_path)

    print(f"🎉 Đã hoàn tất tạo {len(generated_files)} tệp .docx thành công!")
    return generated_files


if __name__ == "__main__":
    asyncio.run(generate_all_86_docx())
