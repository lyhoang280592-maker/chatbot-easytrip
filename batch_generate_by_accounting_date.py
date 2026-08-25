import os
import re
import json
import asyncio
import shutil
from datetime import datetime
import httpx
import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from dotenv import load_dotenv
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from PIL import Image, ImageOps
from generate_contracts import build_contract_pdf

load_dotenv()

LARK_APP_ID = os.getenv("LARK_APP_ID")
LARK_APP_SECRET = os.getenv("LARK_APP_SECRET")
LARK_APP_TOKEN = os.getenv("LARK_APP_TOKEN")
TABLE_ID = "tbluosVi3sQS9gIS"


def num_to_vietnamese_words(n: int) -> str:
    """Chuyển số tiền thành chữ Tiếng Việt chuẩn xác"""
    if n == 0: return "Không đồng"
    don_vi = ["", "nghìn", "triệu", "tỷ", "nghìn tỷ", "triệu tỷ"]
    chu_so = ["không", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
    
    def doc_3_so(num, day_du=False):
        t = num // 100
        c = (num % 100) // 10
        d = num % 10
        res = []
        if t > 0 or day_du:
            res.append(chu_so[t] + " trăm")
        if c > 1:
            res.append(chu_so[c] + " mươi")
            if d == 1: res.append("mốt")
            elif d == 4: res.append("tư")
            elif d == 5: res.append("lăm")
            elif d > 0: res.append(chu_so[d])
        elif c == 1:
            res.append("mười")
            if d == 5: res.append("lăm")
            elif d > 0: res.append(chu_so[d])
        elif c == 0 and (t > 0 or day_du) and d > 0:
            res.append("lẻ")
            res.append(chu_so[d])
        elif c == 0 and d > 0:
            res.append(chu_so[d])
        return " ".join(res)

    chunks = []
    temp = n
    while temp > 0:
        chunks.append(temp % 1000)
        temp //= 1000
        
    words = []
    for i in range(len(chunks)-1, -1, -1):
        c = chunks[i]
        if c > 0:
            day_du = (i < len(chunks)-1)
            w = doc_3_so(c, day_du)
            if w:
                words.append(w)
                if don_vi[i]:
                    words.append(don_vi[i])
                    
    res_str = " ".join(words).strip()
    return res_str.capitalize() + " đồng chẵn."


def num_to_english_words(n: int) -> str:
    """Chuyển số tiền thành chữ Tiếng Anh"""
    to_19 = ['Zero', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine', 'Ten',
             'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen', 'Seventeen', 'Eighteen', 'Nineteen']
    tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
    denom = ['', 'Thousand', 'Million', 'Billion']
    
    def _convert_nn(val):
        if val < 20: return to_19[val]
        for v, d in enumerate(tens):
            if v >= 2:
                if val < (v + 1) * 10:
                    return d + ('' if val % 10 == 0 else ' ' + to_19[val % 10])
        return ''

    def _convert_nnn(val):
        word = ''
        rem = val % 100
        hun = val // 100
        if hun > 0:
            word = to_19[hun] + ' Hundred'
            if rem > 0: word += ' ' + _convert_nn(rem)
        else:
            word = _convert_nn(rem)
        return word

    if n == 0: return "Zero Vietnamese Dong only."
    word = ''
    i = 0
    while n > 0:
        rem = n % 1000
        if rem > 0:
            w = _convert_nnn(rem)
            word = w + (' ' + denom[i] if denom[i] else '') + (' ' + word if word else '')
        n //= 1000
        i += 1
    return word.strip() + " Vietnamese Dong only."


def extract_and_clean_signature(image_path: str, output_sig_path: str, crop_box=None):
    """
    Trích xuất chữ ký từ ảnh hộ chiếu và tách nền trong suốt
    """
    try:
        img = Image.open(image_path).convert("RGBA")
        width, height = img.size
        
        if not crop_box:
            aspect_ratio = height / width
            if aspect_ratio >= 1.15:
                crop_box = (int(width * 0.05), int(height * 0.44), int(width * 0.52), int(height * 0.70))
            elif aspect_ratio <= 0.85:
                crop_box = (int(width * 0.05), int(height * 0.58), int(width * 0.48), int(height * 0.94))
            else:
                crop_box = (int(width * 0.08), int(height * 0.46), int(width * 0.50), int(height * 0.72))
            
        cropped = img.crop(crop_box)
        gray = ImageOps.grayscale(cropped)
        gray = ImageOps.autocontrast(gray, cutoff=2)
        
        datas = gray.getdata()
        new_data = []
        for item in datas:
            val = int(item) if isinstance(item, (int, float)) else int(item[0])  # type: ignore
            if val < 135:
                alpha = int((135 - val) / 135 * 255)
                new_data.append((25, 30, 45, min(255, alpha * 2)))
            else:
                new_data.append((255, 255, 255, 0))
                
        sig_img = Image.new("RGBA", cropped.size)
        sig_img.putdata(new_data)
        
        bbox = sig_img.getbbox()
        if bbox:
            sig_img = sig_img.crop(bbox)
            
        os.makedirs(os.path.dirname(output_sig_path) or ".", exist_ok=True)
        sig_img.save(output_sig_path, "PNG")
        return True
    except Exception as e:
        print(f"⚠️ Error extracting signature from {image_path}: {e}")
        return False



def set_cell_border(cell, **kwargs):
    """Set cell borders in python-docx"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell margins (padding) in twips"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_p(doc, text: str = "", font_name: str = "Arial", font_size: float | int = 10, bold: bool = False, italic: bool = False, 
          align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, space_before: float | int = 0, space_after: float | int = 2, line_spacing: float | int = 1.15):
    """Helper to add styled paragraph"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def build_contract_docx(output_docx_path: str, data: dict):
    """
    Tạo tệp Word .docx Hợp đồng 7 trang chuẩn xác 100% theo mẫu
    """
    doc = docx.Document()
    
    # Thiết lập căn lề chuẩn A4 2.0 cm
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)

    # -------------------------------------------------------------
    # TRANG 1: TIÊU NGỮ & BÊN A
    # -------------------------------------------------------------
    add_p(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "Độc lập - Tự do - Hạnh phúc", font_size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "---------------", font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    add_p(doc, "HỢP ĐỒNG CUNG CẤP DỊCH VỤ TƯ VẤN VÀ LÀM HỒ SƠ VISA", font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc, "CONSULTING AND VISA APPLICATION SERVICE CONTRACT", font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, f"Số / No: {data.get('contract_no', '......')}/2026/HĐDV-EASYTRIP", font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    legal_bases = [
        ("Căn cứ Bộ luật Dân sự số 91/2015/QH13 ngày 24/11/2015 và các văn bản hướng dẫn thi hành;",
         "Pursuant to the Civil Code No. 91/2015/QH13 dated November 24, 2015 and its guiding documents;"),
        ("Căn cứ Luật Thương mại số 36/2005/QH11 ngày 14/06/2005 và các văn bản hướng dẫn thi hành;",
         "Pursuant to the Commercial Law No. 36/2005/QH11 dated June 14, 2005 and its guiding documents;"),
        ("Căn cứ Luật Doanh nghiệp số 59/2020/QH14 ngày 17/06/2020 và các văn bản hướng dẫn thi hành;",
         "Pursuant to the Law on Enterprises No. 59/2020/QH14 dated June 17, 2020 and its guiding documents;"),
        ("Căn cứ Luật Nhập cảnh, xuất cảnh, quá cảnh, cư trú của người nước ngoài tại Việt Nam số 47/2014/QH13 ngày 16/06/2014 và các luật sửa đổi, bổ sung liên quan;",
         "Pursuant to the Law on Entry, Exit, Transit, and Residence of Foreigners in Vietnam No. 47/2014/QH13 dated June 16, 2014 and its related amendments;"),
        ("Căn cứ Luật Quản lý thuế số 38/2019/QH14 ngày 13/06/2019 và Luật Thuế giá trị gia tăng số 13/2008/QH12 cùng các văn bản hướng dẫn thi hành;",
         "Pursuant to the Law on Tax Administration No. 38/2019/QH14 dated June 13, 2019 and the Law on Value Added Tax No. 13/2008/QH12 and their guiding documents;"),
        ("Căn cứ Thông tư số 28/2026/TT-BTC ngày 27/03/2026 quy định về phí và lệ phí trong lĩnh vực xuất cảnh, nhập cảnh, quá cảnh, cư trú tại Việt Nam;",
         "Pursuant to Circular No. 28/2026/TT-BTC dated March 27, 2026 regulating fees and charges in the field of entry, exit, transit, and residence in Vietnam;"),
        ("Căn cứ vào nhu cầu và khả năng thực tế của hai Bên.",
         "Pursuant to the demands and actual capacities of both Parties.")
    ]

    for vi, en in legal_bases:
        add_p(doc, f"- {vi}", font_size=9, space_after=1)
        add_p(doc, en, font_size=8.5, italic=True, space_after=4)

    contract_date_vi = data.get('date_vi', 'ngày 16 tháng 08 năm 2026')
    contract_date_en = data.get('date_en', 'August 16, 2026')

    add_p(doc, f"Hôm nay, {contract_date_vi}, tại văn phòng CÔNG TY TNHH CHUYẾN ĐI VÀ THỊ THỰC DỄ DÀNG, chúng tôi gồm có:", font_size=9.5, space_before=4, space_after=1)
    add_p(doc, f"Today, {contract_date_en}, at the office of EASY TRIP AND VISA COMPANY LIMITED, we consist of:", font_size=9, italic=True, space_after=6)

    add_p(doc, "BÊN A: BÊN SỬ DỤNG DỊCH VỤ (KHÁCH HÀNG)", font_size=9.5, bold=True, space_after=1)
    add_p(doc, "PARTY A: SERVICE USER (CUSTOMER)", font_size=9.5, bold=True, space_after=4)

    # Customer info
    p_name = add_p(doc, font_size=9.5, space_after=2)
    p_name.add_run("Tên tổ chức/cá nhân / Name: ")
    r_name = p_name.add_run(str(data.get('customer_name', '')).upper())
    r_name.bold = True

    p_pass = add_p(doc, font_size=9.5, space_after=2)
    p_pass.add_run("Số Hộ chiếu / Passport No.: ")
    r_pass = p_pass.add_run(str(data.get('passport_no', '')))
    r_pass.bold = True
    p_pass.add_run("    Ngày cấp / Date of Issue: ")
    r_doi = p_pass.add_run(str(data.get('date_of_issue', '20/05/2022')))
    r_doi.bold = True

    p_nat = add_p(doc, font_size=9.5, space_after=2)
    p_nat.add_run("Quốc tịch / Nationality: ")
    r_nat = p_nat.add_run(str(data.get('nationality', 'Nga / Russian')))
    r_nat.bold = True

    add_p(doc, f"Email: {data.get('email', '.................................................................................')}", font_size=9.5, space_after=2)
    add_p(doc, "Đại diện bởi / Represented by: .......................................................", font_size=9.5, space_after=4)

    # -------------------------------------------------------------
    # TRANG 2: BÊN B & ĐIỀU 1
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "BÊN B: BÊN CUNG CẤP DỊCH VỤ", font_size=9.5, bold=True, space_after=1)
    add_p(doc, "PARTY B: SERVICE PROVIDER", font_size=9.5, bold=True, space_after=4)

    p_b_name = add_p(doc, font_size=9.5, space_after=1)
    p_b_name.add_run("Tên đơn vị / Company Name: ")
    r_bn = p_b_name.add_run("CÔNG TY TNHH CHUYẾN ĐI VÀ THỊ THỰC DỄ DÀNG")
    r_bn.bold = True

    add_p(doc, "English Name: EASY TRIP AND VISA COMPANY LIMITED", font_size=9, italic=True, space_after=2)
    add_p(doc, "Mã số thuế / Tax Code: 4202051389", font_size=9.5, space_after=2)
    add_p(doc, "Địa chỉ / Address: 21 Phan Vinh, Phường Vĩnh Nguyên, TP. Nha Trang, tỉnh Khánh Hòa, Việt Nam", font_size=9.5, space_after=2)

    p_rep = add_p(doc, font_size=9.5, space_after=2)
    p_rep.add_run("Đại diện / Represented by: Ông/Bà ")
    r_rep = p_rep.add_run("LÝ VIỆT HOÀNG")
    r_rep.bold = True
    p_rep.add_run(" (Mr./Ms. LY VIET HOANG)")

    add_p(doc, "Chức vụ / Position: Giám đốc / Director", font_size=9.5, space_after=2)
    add_p(doc, "Điện thoại / Telephone: 0896916361 - Hotline: 0896916361", font_size=9.5, space_after=2)

    p_acc = add_p(doc, font_size=9.5, space_after=4)
    p_acc.add_run("Tài khoản ngân hàng / Bank Account: VCB Khánh Hòa Số Tài khoản / Account No.: ")
    r_acc = p_acc.add_run("1068582577")
    r_acc.bold = True

    add_p(doc, "Hai Bên cùng thống nhất ký kết Hợp đồng dịch vụ làm thủ tục xin Visa với các điều khoản sau:", font_size=9.5, space_before=4, space_after=1)
    add_p(doc, "Both Parties hereby agree to enter into this Visa Service Contract with the following terms and conditions:", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 1: NỘI DUNG DỊCH VỤ", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 1: SCOPE OF SERVICES", font_size=10, bold=True, space_after=4)

    add_p(doc, "1.1 Bên A giao và Bên B đồng ý thực hiện dịch vụ tư vấn, chuẩn bị, hoàn thiện bộ hồ sơ để nộp hồ sơ và nộp lệ phí xin cấp Visa (Thị thực) cho Bên A với chi tiết sau:", font_size=9.5, space_after=1)
    add_p(doc, "1.1 Party A delegates and Party B agrees to perform the services of consulting preparing, completing, and payment of visa submitting visa applications for Party A as detailed below:", font_size=9, italic=True, space_after=6)

    # Table 1: Scope of Services
    t1 = doc.add_table(rows=2, cols=6)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths1 = [Cm(1.0), Cm(3.8), Cm(2.5), Cm(4.2), Cm(2.2), Cm(2.8)]

    headers1 = [
        ("STT\nNo.", True),
        ("Họ và tên người xin Visa\nFull name of applicant", True),
        ("Số Hộ chiếu\nPassport No.", True),
        ("Cơ quan cấp thị thực\nIssuing Authority", True),
        ("Mục đích\nPurpose", True),
        ("Loại Visa\nVisa Type", True)
    ]

    for c_idx, (h_text, is_bold) in enumerate(headers1):
        cell = t1.cell(0, c_idx)
        cell.width = col_widths1[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = h_text.split("\n")
        r1 = p.add_run(lines[0])
        r1.font.name = "Arial"
        r1.font.size = Pt(8.5)
        r1.font.bold = True
        if len(lines) > 1:
            p.add_run("\n")
            r2 = p.add_run(lines[1])
            r2.font.name = "Arial"
            r2.font.size = Pt(8.0)
            r2.font.italic = True

    row1_data = [
        "1",
        str(data.get('customer_name', '')).upper(),
        str(data.get('passport_no', '')),
        "Cục QL XNK - Bộ Công an\nImmigration Dept - MPS (Vietnam)",
        "Du lịch\nTourism",
        "Thị thực điện tử /\nEVisa"
    ]

    for c_idx, val in enumerate(row1_data):
        cell = t1.cell(1, c_idx)
        cell.width = col_widths1[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = val.split("\n")
        r1 = p.add_run(lines[0])
        r1.font.name = "Arial"
        r1.font.size = Pt(8.5)
        if c_idx == 1:
            r1.font.bold = True
        if len(lines) > 1:
            p.add_run("\n")
            r2 = p.add_run(lines[1])
            r2.font.name = "Arial"
            r2.font.size = Pt(8.0)
            r2.font.italic = True

    add_p(doc, space_after=6)
    add_p(doc, "1.2 Nội dung công việc chi tiết của Bên B:", font_size=9.5, space_after=1)
    add_p(doc, "1.2 Detailed scope of work of Party B:", font_size=9, italic=True, space_after=4)
    add_p(doc, "- Tư vấn các quy định pháp lý, điều kiện và thủ tục xin cấp thị thực điện tử của Cục Quản lý Xuất nhập cảnh - Bộ Công an.", font_size=9.5, space_after=1)
    add_p(doc, "Consult on legal regulations, conditions, and procedures for e-visa issuance of the Immigration Department - Ministry of Public Security.", font_size=9, italic=True, space_after=4)
    add_p(doc, "- Hướng dẫn Bên A chuẩn bị giấy tờ, thông tin cá nhân đúng tiêu chuẩn theo yêu cầu.", font_size=9.5, space_after=1)

    # -------------------------------------------------------------
    # TRANG 3: ĐIỀU 2 (CHI PHÍ & THANH TOÁN)
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "Guide Party A to prepare documents and personal information up to standard as required.", font_size=9, italic=True, space_after=4)
    add_p(doc, "- Thực hiện khai trực tuyến thông tin, điền tờ khai và nộp hồ sơ đăng ký.", font_size=9.5, space_after=1)
    add_p(doc, "Perform online declaration, fill out application forms, and submit registration documents.", font_size=9, italic=True, space_after=4)
    add_p(doc, "- Theo dõi tiến trình xử lý của Bộ Công an và bàn giao kết quả thị thực điện tử cho Bên A ngay sau khi có kết quả.", font_size=9.5, space_after=1)
    add_p(doc, "Track the processing status by the Ministry of Public Security and hand over the electronic visa results to Party A immediately upon availability.", font_size=9, italic=True, space_after=8)

    add_p(doc, "ĐIỀU 2: GIÁ TRỊ HỢP ĐỒNG VÀ PHƯƠNG THỨC THANH TOÁN", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 2: CONTRACT VALUE AND PAYMENT METHOD", font_size=10, bold=True, space_after=4)
    add_p(doc, "2.1 Bảng chi tiết chi phí / Detailed Cost Breakdown:", font_size=9.5, space_after=6)

    total_amount = data.get('total_amount')
    state_fee = int(data.get('state_fee', 662500))
    transport_fee = int(data.get('transport_fee', 0))

    if total_amount is not None:
        combined_service_fee = int(total_amount) - state_fee
        pure_consulting_fee = int(data.get('service_fee', combined_service_fee - transport_fee))
    else:
        pure_consulting_fee = int(data.get('service_fee', 1637500))
        combined_service_fee = pure_consulting_fee + transport_fee
        total_amount = combined_service_fee + state_fee

    # Table 2: Cost breakdown
    t2 = doc.add_table(rows=4, cols=5)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths2 = [Cm(1.0), Cm(7.5), Cm(1.6), Cm(3.2), Cm(3.2)]

    headers2 = [
        "STT\nNo.",
        "Nội dung chi phí\nDescription of Fees",
        "Số lượng\nQty",
        "Đơn giá (VNĐ)\nUnit Price (VND)",
        "Thành tiền (VNĐ)\nTotal Amount (VND)"
    ]

    for c_idx, h_text in enumerate(headers2):
        cell = t2.cell(0, c_idx)
        cell.width = col_widths2[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = h_text.split("\n")
        r1 = p.add_run(lines[0])
        r1.font.name = "Arial"
        r1.font.size = Pt(8.5)
        r1.font.bold = True
        if len(lines) > 1:
            p.add_run("\n")
            r2 = p.add_run(lines[1])
            r2.font.name = "Arial"
            r2.font.size = Pt(8.0)
            r2.font.italic = True

    # Row 1: Phí dịch vụ gộp
    r1_items = [
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        ("Phí dịch vụ tư vấn, làm hồ sơ Visa và hỗ trợ vận tải\nService fee for consulting, visa processing & passenger transport support", WD_ALIGN_PARAGRAPH.LEFT),
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        (f"{combined_service_fee:,}".replace(",", "."), WD_ALIGN_PARAGRAPH.CENTER),
        (f"{combined_service_fee:,}".replace(",", "."), WD_ALIGN_PARAGRAPH.CENTER)
    ]
    for c_idx, (text_val, align_val) in enumerate(r1_items):
        cell = t2.cell(1, c_idx)
        cell.width = col_widths2[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = align_val
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = text_val.split("\n")
        r = p.add_run(lines[0])
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        if len(lines) > 1:
            p.add_run("\n")
            r_sub = p.add_run(lines[1])
            r_sub.font.name = "Arial"
            r_sub.font.size = Pt(8.0)
            r_sub.font.italic = True

    # Row 2: Lệ phí nhà nước
    usd_val = int(round(state_fee / 26500))
    if usd_val <= 0: usd_val = 25
    r2_desc = f"Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử (Thu hộ - Chi hộ)\n(${usd_val} – tỷ giá 26.500VND)\nState fee for electronic visa registration (Collected & Paid on behalf)\n(${usd_val} – exchange rate 26,500 VND)"

    r2_items = [
        ("2", WD_ALIGN_PARAGRAPH.CENTER),
        (r2_desc, WD_ALIGN_PARAGRAPH.LEFT),
        ("1", WD_ALIGN_PARAGRAPH.CENTER),
        (f"{state_fee:,}".replace(",", "."), WD_ALIGN_PARAGRAPH.CENTER),
        (f"{state_fee:,}".replace(",", "."), WD_ALIGN_PARAGRAPH.CENTER)
    ]
    for c_idx, (text_val, align_val) in enumerate(r2_items):
        cell = t2.cell(2, c_idx)
        cell.width = col_widths2[c_idx]
        set_cell_border(cell)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = align_val
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lines = text_val.split("\n")
        for l_idx, line in enumerate(lines):
            if l_idx > 0: p.add_run("\n")
            r = p.add_run(line)
            r.font.name = "Arial"
            r.font.size = Pt(8.5 if l_idx < 2 else 8.0)
            if l_idx >= 2: r.font.italic = True

    # Row 3: Merge 0..3 and Total
    cell_tot_lbl = t2.cell(3, 0)
    cell_tot_lbl.merge(t2.cell(3, 3))
    set_cell_border(cell_tot_lbl)
    set_cell_margins(cell_tot_lbl, 120, 120, 150, 150)
    p_tot_lbl = cell_tot_lbl.paragraphs[0]
    p_tot_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tot_lbl.paragraph_format.space_before = Pt(0)
    p_tot_lbl.paragraph_format.space_after = Pt(0)
    r_tl = p_tot_lbl.add_run("Tổng cộng / Total Amount")
    r_tl.font.name = "Arial"
    r_tl.font.size = Pt(8.5)
    r_tl.font.bold = True

    cell_tot_val = t2.cell(3, 4)
    cell_tot_val.width = col_widths2[4]
    set_cell_border(cell_tot_val)
    set_cell_margins(cell_tot_val, 120, 120, 150, 150)
    p_tot_val = cell_tot_val.paragraphs[0]
    p_tot_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tot_val.paragraph_format.space_before = Pt(0)
    p_tot_val.paragraph_format.space_after = Pt(0)
    r_tv = p_tot_val.add_run(f"{total_amount:,}".replace(",", "."))
    r_tv.font.name = "Arial"
    r_tv.font.size = Pt(8.5)
    r_tv.font.bold = True

    add_p(doc, space_after=4)

    words_vi = num_to_vietnamese_words(total_amount)
    words_en = num_to_english_words(total_amount)

    p_wvi = add_p(doc, font_size=9.5, space_after=1)
    p_wvi.add_run("Bằng chữ: ")
    r_wvi = p_wvi.add_run(words_vi)
    r_wvi.bold = True

    p_wen = add_p(doc, font_size=9, italic=True, space_after=6)
    p_wen.add_run("In words: ")
    p_wen.add_run(words_en)

    add_p(doc, "2.2 Thuế Giá trị gia tăng (VAT) / Value Added Tax (VAT):", font_size=9.5, bold=True, space_after=2)
    add_p(doc, "- Đơn giá phí dịch vụ nêu trên đã bao gồm thuế GTGT theo quy định của pháp luật Việt Nam. Bên B có trách nhiệm lập và xuất hóa đơn điện tử giá trị gia tăng (GTGT) hợp pháp cho Bên A đối với phần Phí dịch vụ tư vấn, làm hồ sơ Visa và hỗ trợ vận tải trong vòng 03 ngày làm việc kể từ ngày hoàn thành dịch vụ hoặc khi Bên A thanh toán đầy đủ, tùy điều kiện nào đến trước. Bên A có trách nhiệm cung cấp đầy đủ và chính xác thông tin xuất hóa đơn.", font_size=9.5, space_after=2)
    add_p(doc, "- The service fee specified above is inclusive of VAT in accordance with Vietnamese tax laws. Party B is responsible for issuing a valid electronic Value Added Tax (VAT) invoice to Party A for the consulting, visa processing & transport support service fee component within 03 working days from the completion of the service or upon full payment by Party A, whichever comes first. Party A is responsible for providing complete and accurate billing information.", font_size=9, italic=True, space_after=4)

    # -------------------------------------------------------------
    # TRANG 4: ĐIỀU 3 & ĐIỀU 4
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "- Đối với phần Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử (Thu hộ - Chi hộ), Bên B sẽ cung cấp chứng từ, biên lai lệ phí của Bộ Công an/Cổng dịch vụ công cho Bên A và phần này không chịu thuế GTGT của Bên B.", font_size=9.5, space_after=2)
    add_p(doc, "- For the State fee for electronic visa registration (Collected & Paid on behalf), Party B shall provide the official receipt from the Ministry of Public Security / Public Service Portal to Party A, and this component is not subject to VAT by Party B.", font_size=9, italic=True, space_after=4)

    add_p(doc, "2.3 Tiến độ thanh toán / Payment Schedule:", font_size=9.5, bold=True, space_after=1)
    add_p(doc, "- Bên A thanh toán 100% tổng giá trị hợp đồng ngay sau khi ký kết Hợp đồng này.", font_size=9.5, space_after=1)
    add_p(doc, "- Party A shall pay 100% of the total contract value immediately upon signing this Contract.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 3: QUYỀN VÀ NGHĨA VỤ CỦA BÊN A", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 3: RIGHTS AND OBLIGATIONS OF PARTY A", font_size=10, bold=True, space_after=4)

    add_p(doc, "3.1 Cung cấp đầy đủ, trung thực, chính xác và kịp thời các hồ sơ, thông tin, giấy tờ theo hướng dẫn của Bên B. Bên A tự chịu hoàn toàn trách nhiệm pháp lý trước pháp luật về tính chân thật, hợp pháp của các tài liệu đã cung cấp.", font_size=9.5, space_after=1)
    add_p(doc, "3.1 Provide fully, honestly, accurately, and timely all dossiers, information, and documents as guided by Party B. Party A shall be solely and fully liable under the law for the authenticity and legality of the provided documents.", font_size=9, italic=True, space_after=4)

    add_p(doc, "3.2 Thanh toán đầy đủ và đúng hạn các khoản phí theo quy định tại Điều 2 của Hợp đồng này.", font_size=9.5, space_after=1)
    add_p(doc, "3.2 Pay fully and on time all fees specified in Article 2 of this Contract.", font_size=9, italic=True, space_after=4)

    add_p(doc, "3.3 Có mặt đúng giờ hoặc phối hợp cung cấp sinh trắc học, thông tin bổ sung khi có yêu cầu từ Cục Quản lý Xuất nhập cảnh - Bộ Công an hoặc cơ quan có thẩm quyền của Nhà nước.", font_size=9.5, space_after=1)
    add_p(doc, "3.3 Be present on time or coordinate to provide biometrics and additional information upon request from the Immigration Department - Ministry of Public Security or competent State authorities.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 4: QUYỀN VÀ NGHĨA VỤ CỦA BÊN B", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 4: RIGHTS AND OBLIGATIONS OF PARTY B", font_size=10, bold=True, space_after=4)

    add_p(doc, "4.1 Được quyền sử dụng thông tin trên hộ chiếu, thông tin khác của bên A để thực hiện công việc tư vấn và chuẩn bị hồ sơ đăng ký các dịch vụ công của Nhà Nước Việt Nam.", font_size=9.5, space_after=1)
    add_p(doc, "4.1 To be authorized to use Party A's passport information and other information to perform consultancy services and prepare application dossiers for public services of the State of Vietnam.", font_size=9, italic=True, space_after=4)

    add_p(doc, "4.2 Bảo quản cẩn thận, an toàn các giấy tờ gốc, thông tin cá nhân do Bên A bàn giao trong quá trình thực hiện hợp đồng.", font_size=9.5, space_after=1)
    add_p(doc, "4.2 Carefully and safely preserve the original documents and personal information handed over by Party A during the contract execution.", font_size=9, italic=True, space_after=4)

    add_p(doc, "4.3 Bảo mật tuyệt đối mọi thông tin cá nhân, thông tin hồ sơ của Bên A và không tiết lộ cho bất kỳ bên thứ ba nào khi chưa có sự đồng ý bằng văn bản của Bên A, trừ trường hợp cung cấp cho Cục Quản lý Xuất nhập cảnh - Bộ Công an để đăng ký cấp thị thực điện tử hoặc theo yêu cầu của pháp luật.", font_size=9.5, space_after=1)
    add_p(doc, "4.3 Keep strictly confidential all personal and dossier information of Party A and not disclose it to any third party without Party A's prior written consent, except", font_size=9, italic=True, space_after=4)

    # -------------------------------------------------------------
    # TRANG 5: ĐIỀU 5 & ĐIỀU 6
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "for submission to the Immigration Department - Ministry of Public Security for e-visa registration or as required by law.", font_size=9, italic=True, space_after=4)
    add_p(doc, "4.4 Thông báo kịp thời cho Bên A về tiến độ và kết quả hồ sơ. Bàn giao đầy đủ kết quả dịch vụ công ngay sau khi được Nhà Nước Việt Nam ban hành.", font_size=9.5, space_after=1)
    add_p(doc, "4.4 Promptly notify Party A of the progress and results of the application. Fully hand over the results of the public service immediately upon issuance by the State of Vietnam.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 5: ĐIỀU KHOẢN VỀ KẾT QUẢ DỊCH VỤ VÀ HOÀN PHÍ", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 5: SERVICE RESULTS AND REFUND POLICY", font_size=10, bold=True, space_after=4)

    add_p(doc, "5.1 Hai Bên hiểu rõ rằng quyền quyết định cấp hoặc từ chối cấp thị thực điện tử hoàn toàn thuộc thẩm quyền của Cục Quản lý Xuất nhập cảnh - Bộ Công an nước Cộng hòa Xã hội Chủ nghĩa Việt Nam. Bên B không quyết định và không bảo đảm tuyệt đối kết quả cấp thị thực.", font_size=9.5, space_after=1)
    add_p(doc, "5.1 Both Parties clearly understand that the decision to grant or refuse the electronic visa belongs solely to the authority of the Immigration Department - Ministry of Public Security of the Socialist Republic of Vietnam. Party B does not decide and does not guarantee absolute visa issuance results.", font_size=9, italic=True, space_after=4)

    add_p(doc, "5.2 Xử lý khi rớt thị thực do lỗi khách quan (từ phía Cục Quản lý Xuất nhập cảnh - Bộ Công an từ chối mà không do lỗi của Bên nào):", font_size=9.5, space_after=1)
    add_p(doc, "5.2 Handling visa rejection due to objective reasons (rejection by the Immigration Department - Ministry of Public Security without fault of either Party):", font_size=9, italic=True, space_after=4)

    add_p(doc, "- Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử sẽ không được hoàn lại (theo quy định của Bộ Công an và Thông tư số 28/2026/TT-BTC).", font_size=9.5, space_after=1)
    add_p(doc, "The State fee for electronic visa registration will not be refunded (in accordance with the Ministry of Public Security regulations and Circular No. 28/2026/TT-BTC).", font_size=9, italic=True, space_after=4)

    add_p(doc, f"- Bên B sẽ hoàn trả lại 100% Phí dịch vụ tư vấn & làm hồ sơ ({pure_consulting_fee:,} VNĐ) cho Bên A trong vòng 05 ngày làm việc kể từ ngày nhận được thông báo từ chối cấp thị thực (không bao gồm Lệ phí nộp Nhà nước và Phí dịch vụ hỗ trợ vận tải).".replace(",", "."), font_size=9.5, space_after=1)
    add_p(doc, f"Party B shall refund 100% of the Visa Consulting & Processing Service Fee (VND {pure_consulting_fee:,}) to Party A within 05 working days from the date of receiving the visa rejection notice (excluding the State fee and Passenger transport support fee).".replace(",", "."), font_size=9, italic=True, space_after=4)

    add_p(doc, "5.3 Trường hợp Bên A đơn phương hủy hợp đồng sau khi Bên B đã tiến hành xử lý hồ sơ hoặc nộp lệ phí, phí dịch vụ và Lệ phí nộp Nhà nước đăng ký cấp thị thực điện tử sẽ không được hoàn lại.", font_size=9.5, space_after=1)
    add_p(doc, "5.3 In case Party A unilaterally terminates the contract after Party B has commenced processing the dossier or paid the fees, the service fee and the State fee for electronic visa registration shall not be refunded.", font_size=9, italic=True, space_after=4)

    add_p(doc, "5.4 Trường hợp Bên A cung cấp thông tin, tài liệu giả mạo hoặc sai sự thật dẫn đến việc hồ sơ bị từ chối hoặc bị xử lý theo pháp luật:", font_size=9.5, space_after=1)
    add_p(doc, "5.4 In case Party A provides forged or untruthful information or documents, leading to visa rejection or legal actions:", font_size=9, italic=True, space_after=4)

    add_p(doc, "- Bên B có quyền đơn phương chấm dứt hợp đồng ngay lập tức.", font_size=9.5, space_after=1)
    add_p(doc, "Party B has the right to unilaterally terminate the contract immediately.", font_size=9, italic=True, space_after=4)

    add_p(doc, "- Bên A không được hoàn lại bất kỳ khoản phí nào và phải tự chịu hoàn toàn trách nhiệm trước pháp luật.", font_size=9.5, space_after=1)
    add_p(doc, "Party A shall not be refunded any fees and must bear full legal responsibility.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 6: BẤT KHẢ KHÁNG VÀ GIỚI HẠN TRÁCH NHIỆM", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 6: FORCE MAJEURE AND LIMITATION OF LIABILITY", font_size=10, bold=True, space_after=4)

    # -------------------------------------------------------------
    # TRANG 6: ĐIỀU 7 & ĐIỀU 8
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "6.1 Sự kiện bất khả kháng là sự kiện xảy ra một cách khách quan, không thể lường trước được và không thể khắc phục được mặc dù đã áp dụng mọi biện pháp cần thiết và khả năng cho phép, bao gồm nhưng không giới hạn ở: thiên tai, dịch bệnh, chiến tranh, sự thay đổi đột ngột về chính sách xuất nhập cảnh của Chính phủ Việt Nam hoặc quốc gia liên quan, việc Cục Quản lý Xuất nhập cảnh - Bộ Công an tạm ngừng hoạt động hoặc đóng cửa biên giới.", font_size=9.5, space_after=1)
    add_p(doc, "6.1 A Force Majeure event is an event that occurs objectively, unpredictably, and irremediably despite all necessary and permissible measures being taken, including but not limited to: natural disasters, epidemics, wars, sudden changes in immigration policies of the Vietnamese Government or related countries, temporary suspension of operations of the Immigration Department - Ministry of Public Security, or border closures.", font_size=9, italic=True, space_after=4)

    add_p(doc, "6.2 Trong trường hợp xảy ra Sự kiện bất khả kháng dẫn đến việc chậm trễ hoặc không thể thực hiện nghĩa vụ hợp đồng, bên bị ảnh hưởng sẽ được miễn trừ trách nhiệm và không phải bồi thường thiệt hại, với điều kiện phải thông báo cho bên kia bằng văn bản trong vòng 03 ngày kể từ ngày xảy ra sự kiện. Hai bên sẽ cùng thương lượng để tìm giải pháp khắc phục.", font_size=9.5, space_after=1)
    add_p(doc, "6.2 In the event of a Force Majeure event leading to delay or inability to perform contractual obligations, the affected party shall be exempted from liability and shall not compensate for damages, provided that the other party is notified in writing within 03 days from the occurrence of the event. Both parties shall negotiate to find a remedy.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 7: GIẢI QUYẾT TRANH CHẤP", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 7: DISPUTE RESOLUTION", font_size=10, bold=True, space_after=4)

    add_p(doc, "7.1 Mọi tranh chấp phát sinh từ hoặc liên quan đến Hợp đồng này trước hết sẽ được hai Bên giải quyết thông qua thương lượng, hòa giải trên tinh thần hợp tác và cùng có lợi.", font_size=9.5, space_after=1)
    add_p(doc, "7.1 Any dispute arising from or related to this Contract shall first be resolved by both Parties through negotiation and conciliation in a cooperative and mutually beneficial spirit.", font_size=9, italic=True, space_after=4)

    add_p(doc, "7.2 Trường hợp tranh chấp không thể giải quyết bằng thương lượng trong vòng 30 ngày kể từ ngày phát sinh, một trong hai Bên có quyền đưa vụ việc ra giải quyết tại Tòa án nhân dân có thẩm quyền tại Thành phố Nha Trang, tỉnh Khánh Hòa, Việt Nam để giải quyết theo quy định của pháp luật Việt Nam.", font_size=9.5, space_after=1)
    add_p(doc, "7.2 In case the dispute cannot be resolved through negotiation within 30 days from its occurrence, either Party has the right to refer the dispute to the competent People's Court in Nha Trang City, Khanh Hoa Province, Vietnam for resolution in accordance with Vietnamese law.", font_size=9, italic=True, space_after=6)

    add_p(doc, "ĐIỀU 8: ĐIỀU KHOẢN CHUNG", font_size=10, bold=True, space_after=1)
    add_p(doc, "ARTICLE 8: GENERAL PROVISIONS", font_size=10, bold=True, space_after=4)

    add_p(doc, "8.1 Hợp đồng này được điều chỉnh và giải thích theo quy định của pháp luật nước Cộng hòa Xã hội Chủ nghĩa Việt Nam.", font_size=9.5, space_after=1)
    add_p(doc, "8.1 This Contract shall be governed by and construed in accordance with the laws of the Socialist Republic of Vietnam.", font_size=9, italic=True, space_after=4)

    add_p(doc, "8.2 Hợp đồng này có hiệu lực kể từ ngày ký và tự động thanh lý sau khi Bên B bàn giao kết quả thị thực điện tử cho Bên A và Bên A hoàn tất nghĩa vụ thanh toán.", font_size=9.5, space_after=1)
    add_p(doc, "8.2 This Contract shall take effect from the date of signing and shall be automatically liquidated after Party B hands over the electronic visa results to Party A and Party A completes the payment obligations.", font_size=9, italic=True, space_after=6)

    # -------------------------------------------------------------
    # TRANG 7: KÝ TÊN & ĐÓNG DẤU
    # -------------------------------------------------------------
    doc.add_page_break()

    add_p(doc, "8.3 Hợp đồng được lập thành 02 (hai) bản bằng tiếng Việt và tiếng Anh. Mỗi bên giữ 01 (một) bản có giá trị pháp lý như nhau. Bản tiếng Việt sẽ được ưu tiên áp dụng nếu có bất kỳ sự khác biệt nào về cách giải thích giữa hai ngôn ngữ.", font_size=9.5, space_after=1)
    add_p(doc, "8.3 The Contract is prepared in 02 (two) copies in both Vietnamese and English. Each party shall keep 01 (one) copy with equal legal validity. The Vietnamese version shall prevail in case of any discrepancy in interpretation between the two languages.", font_size=9, italic=True, space_after=20)

    # Bảng chữ ký 2 cột
    t_sig = doc.add_table(rows=3, cols=2)
    t_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths_sig = [Cm(8.5), Cm(8.5)]

    # Header row
    c_a_hdr = t_sig.cell(0, 0)
    c_a_hdr.width = col_widths_sig[0]
    p_ah = c_a_hdr.paragraphs[0]
    p_ah.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ah1 = p_ah.add_run("ĐẠI DIỆN BÊN A\n")
    r_ah1.bold = True
    r_ah1.font.name = "Arial"
    r_ah1.font.size = Pt(9.5)
    r_ah2 = p_ah.add_run("REPRESENTATIVE OF PARTY A")
    r_ah2.italic = True
    r_ah2.font.name = "Arial"
    r_ah2.font.size = Pt(8.5)

    c_b_hdr = t_sig.cell(0, 1)
    c_b_hdr.width = col_widths_sig[1]
    p_bh = c_b_hdr.paragraphs[0]
    p_bh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bh1 = p_bh.add_run("ĐẠI DIỆN BÊN B\n")
    r_bh1.bold = True
    r_bh1.font.name = "Arial"
    r_bh1.font.size = Pt(9.5)
    r_bh2 = p_bh.add_run("CÔNG TY TNHH CHUYẾN ĐI VÀ THỊ THỰC DỄ DÀNG\n")
    r_bh2.bold = True
    r_bh2.font.name = "Arial"
    r_bh2.font.size = Pt(9.5)
    r_bh3 = p_bh.add_run("REPRESENTATIVE OF PARTY B\nEASY TRIP AND VISA CO., LTD\n")
    r_bh3.italic = True
    r_bh3.font.name = "Arial"
    r_bh3.font.size = Pt(8.5)
    r_bh4 = p_bh.add_run("GIÁM ĐỐC\n")
    r_bh4.bold = True
    r_bh4.font.name = "Arial"
    r_bh4.font.size = Pt(9.5)
    r_bh5 = p_bh.add_run("DIRECTOR")
    r_bh5.italic = True
    r_bh5.font.name = "Arial"
    r_bh5.font.size = Pt(8.5)

    # Signature row
    c_a_sig = t_sig.cell(1, 0)
    c_a_sig.width = col_widths_sig[0]
    p_as = c_a_sig.paragraphs[0]
    p_as.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_as.paragraph_format.space_before = Pt(8)
    p_as.paragraph_format.space_after = Pt(8)

    p_as.add_run("\n\n\n")

    c_b_sig = t_sig.cell(1, 1)
    c_b_sig.width = col_widths_sig[1]
    p_bs = c_b_sig.paragraphs[0]
    p_bs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bs.paragraph_format.space_before = Pt(8)
    p_bs.paragraph_format.space_after = Pt(8)
    p_bs.add_run("\n\n\n")

    # Name row
    c_a_name = t_sig.cell(2, 0)
    c_a_name.width = col_widths_sig[0]
    p_an = c_a_name.paragraphs[0]
    p_an.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ann = p_an.add_run(str(data.get('customer_name', '')).upper())
    r_ann.bold = True
    r_ann.font.name = "Arial"
    r_ann.font.size = Pt(9.5)

    c_b_name = t_sig.cell(2, 1)
    c_b_name.width = col_widths_sig[1]
    p_bn = c_b_name.paragraphs[0]
    p_bn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bnn = p_bn.add_run("LÝ VIỆT HOÀNG")
    r_bnn.bold = True
    r_bnn.font.name = "Arial"
    r_bnn.font.size = Pt(9.5)

    os.makedirs(os.path.dirname(output_docx_path) or ".", exist_ok=True)
    doc.save(output_docx_path)
    return output_docx_path


def normalize_nationality(nat_val, name="", code_ev=""):
    if "BUSBY" in name.upper():
        return "Vương Quốc Anh / United Kingdom"
    if nat_val:
        if isinstance(nat_val, list):
            n_str = str(nat_val[0]).strip()
        else:
            n_str = str(nat_val).strip()
        
        n_lower = n_str.lower()
        if n_lower in ["uk", "united kingdom", "anh", "british"]:
            return "Vương Quốc Anh / United Kingdom"
        elif any(k in n_lower for k in ["russia", "nga", "rus"]):
            return "Nga / Russian"
        elif any(k in n_lower for k in ["korea", "hàn", "kor"]):
            return "Hàn Quốc / Korean"
        elif any(k in n_lower for k in ["usa", "mỹ", "american"]):
            return "Mỹ / USA"
        elif any(k in n_lower for k in ["german", "đức", "deu"]):
            return "Đức / Germany"
        elif any(k in n_lower for k in ["australia", "úc", "aus"]):
            return "Úc / Australia"
        elif any(k in n_lower for k in ["turkey", "thổ"]):
            return "Thổ Nhĩ Kỳ / Turkey"
        elif "slovak" in n_lower:
            return "Slovakia / Slovak Republic"
        elif "uzbek" in n_lower:
            return "Uzbekistan"
        elif "ukrain" in n_lower:
            return "Ukraine / Ukrainian"
        elif "kyrgyz" in n_lower:
            return "Kyrgyzstan"
        elif "moldova" in n_lower:
            return "Moldova"
        elif "kazakh" in n_lower:
            return "Kazakhstan"
        elif "brazil" in n_lower:
            return "Brazil"
        elif any(k in n_lower for k in ["netherland", "hà lan", "dutch"]):
            return "Hà Lan / Netherlands"
        elif "belarus" in n_lower:
            return "Belarus"
        elif any(k in n_lower for k in ["france", "pháp"]):
            return "Pháp / France"
        elif any(k in n_lower for k in ["china", "trung"]):
            return "Trung Quốc / China"
        return n_str
    
    # Fallback from code_ev
    if "RUS" in code_ev: return "Nga / Russian"
    if "KOR" in code_ev: return "Hàn Quốc / Korean"
    if "USA" in code_ev: return "Mỹ / USA"
    if "DEU" in code_ev or "GER" in code_ev: return "Đức / Germany"
    if "GBR" in code_ev or "UK" in code_ev: return "Vương Quốc Anh / United Kingdom"
    if "SVK" in code_ev: return "Slovakia / Slovak Republic"
    if "TUR" in code_ev: return "Thổ Nhĩ Kỳ / Turkey"
    if "UZB" in code_ev: return "Uzbekistan"
    if "UKR" in code_ev: return "Ukraine / Ukrainian"
    if "KGZ" in code_ev: return "Kyrgyzstan"
    if "MDA" in code_ev: return "Moldova"
    if "KAZ" in code_ev: return "Kazakhstan"
    if "BRA" in code_ev: return "Brazil"
    if "NLD" in code_ev: return "Hà Lan / Netherlands"
    if "AUS" in code_ev: return "Úc / Australia"
    if "BLR" in code_ev: return "Belarus"
    
    return "Nga / Russian"


def extract_passport_no(c, name, code_ev):
    if "BUSBY" in name.upper():
        return "127294028"
    if code_ev:
        m = re.search(r'[A-Za-z0-9]{7,9}$', code_ev)
        if m:
            return m.group(0)
    ev_files = c.get("EV")
    if ev_files and isinstance(ev_files, list) and len(ev_files) > 0:
        ev_fname = str(ev_files[0].get("name", "") or "")
        m_ev = re.search(r'([A-Za-z0-9]{7,9})\.pdf', ev_fname)
        if m_ev:
            return m_ev.group(1)
        m_num = re.search(r'\d{7,9}', ev_fname)
        if m_num:
            return m_num.group(0)
    return code_ev.replace("E26", "")[:9] if code_ev else "767587433"


async def generate_contracts_by_accounting():
    url = "https://open.larksuite.com/open-apis/auth/v3/tenant_access_token/internal"
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.post(url, json={"app_id": LARK_APP_ID, "app_secret": LARK_APP_SECRET})
        token = r.json().get("tenant_access_token")

    if not token:
        print("❌ Không lấy được Lark Token!")
        return

    headers = {"Authorization": f"Bearer {token}"}
    all_records = []
    page_token = None
    while True:
        p_url = f"https://open.larksuite.com/open-apis/bitable/v1/apps/{LARK_APP_TOKEN}/tables/{TABLE_ID}/records?page_size=500"
        if page_token: p_url += f"&page_token={page_token}"
        async with httpx.AsyncClient(timeout=30) as client:
            r_rec = await client.get(p_url, headers=headers)
            data_res = r_rec.json().get("data", {})
            items = data_res.get("items", [])
            all_records.extend(items)
            if not data_res.get("has_more"): break
            page_token = data_res.get("page_token")

    # MỤC 1. TIÊU CHÍ LỌC DỮ LIỆU TỪ CRM LARK BASE
    start_ts = datetime(2026, 8, 1, 0, 0, 0).timestamp() * 1000
    end_ts = datetime(2026, 8, 24, 23, 59, 59).timestamp() * 1000

    # MỤC 3. QUY TẮC PHÂN LOẠI: Đại lý chỉ gồm đúng 3 đại lý: Bolot, Sergei, Arsenii
    agency_channels = {"bolot", "sergei", "arsenii"}

    matched = []
    seen_customers = set()

    for it in all_records:
        f = it.get("fields", {})
        code_ev = f.get("Code EV")
        ev_file = f.get("EV")
        acc_ts = f.get("Accounting Date")
        name_val = str(f.get("Tên khách (Name)", "") or "").strip()
        note_val = str(f.get("Ghi chú", "") or "").strip()

        # 1. Loại bỏ các hồ sơ bị hủy / cancel
        is_canceled = ("cancel" in name_val.lower()) or ("hủy" in name_val.lower()) or ("huy" in name_val.lower()) or ("cancel" in note_val.lower())
        if is_canceled:
            print(f"🚫 Loại bỏ hồ sơ bị hủy: {name_val}")
            continue

        # 2. Điều kiện lọc: Thời gian + Đã có EV (hoặc hồ sơ Busby)
        if (code_ev or ev_file or "BUSBY" in name_val.upper()) and acc_ts and start_ts <= acc_ts <= end_ts:
            clean_key = re.sub(r'[^a-zA-Z0-9]', '', name_val.upper())
            if clean_key in seen_customers:
                print(f"⚠️ Trùng khách hàng trong tháng 8, loại bỏ bản ghi trùng: {name_val}")
                continue
            seen_customers.add(clean_key)
            matched.append(f)

    # Sắp xếp theo Accounting Date giảm dần
    matched.sort(key=lambda x: x.get("Accounting Date") or 0, reverse=True)
    print(f"📊 Tìm thấy {len(matched)} hồ sơ hợp lệ (đã lọc sạch trùng và hủy) (01/08/2026 - 24/08/2026).")

    # MỤC 6. CẤU TRÚC THƯ MỤC ĐẦU RA VÀ TÀI LIỆU LƯU TRỮ
    base_out_dir = "output_contracts_01_08_den_24_08"

    # Dọn dẹp thư mục cũ để không còn file rác / file hủy
    if os.path.exists(base_out_dir):
        shutil.rmtree(base_out_dir, ignore_errors=True)
    for d_clean in ["output_docx_contracts", "output_contracts"]:
        if os.path.exists(d_clean):
            shutil.rmtree(d_clean, ignore_errors=True)
    
    # Subfolders
    dir_pdf_kl_single = os.path.join(base_out_dir, "PDF", "Khach_Le", "Single")
    dir_pdf_kl_multi = os.path.join(base_out_dir, "PDF", "Khach_Le", "Multi")
    dir_pdf_kl_cam = os.path.join(base_out_dir, "PDF", "Khach_Le", "Visa_Campuchia")
    dir_pdf_dl_single = os.path.join(base_out_dir, "PDF", "Dai_Ly", "Single")
    dir_pdf_dl_multi = os.path.join(base_out_dir, "PDF", "Dai_Ly", "Multi")

    dir_docx_kl_single = os.path.join(base_out_dir, "DOCX", "Khach_Le", "Single")
    dir_docx_kl_multi = os.path.join(base_out_dir, "DOCX", "Khach_Le", "Multi")
    dir_docx_kl_cam = os.path.join(base_out_dir, "DOCX", "Khach_Le", "Visa_Campuchia")
    dir_docx_dl_single = os.path.join(base_out_dir, "DOCX", "Dai_Ly", "Single")
    dir_docx_dl_multi = os.path.join(base_out_dir, "DOCX", "Dai_Ly", "Multi")
    dir_docx_dl_single = os.path.join(base_out_dir, "DOCX", "Dai_Ly", "Single")
    dir_docx_dl_multi = os.path.join(base_out_dir, "DOCX", "Dai_Ly", "Multi")

    for d in [
        dir_pdf_kl_single, dir_pdf_kl_multi, dir_pdf_kl_cam, dir_pdf_dl_single, dir_pdf_dl_multi,
        dir_docx_kl_single, dir_docx_kl_multi, dir_docx_kl_cam, dir_docx_dl_single, dir_docx_dl_multi,
        "output_docx_contracts", "output_contracts", "downloads/passports", "extracted_signatures"
    ]:
        os.makedirs(d, exist_ok=True)

    processed_data = []

    async with httpx.AsyncClient(timeout=30) as client:
        for idx, c in enumerate(matched, 1):
            name = str(c.get("Tên khách (Name)", "CUSTOMER") or "CUSTOMER").strip()
            ch_list = c.get("Nguồn( Channel)", ["Direct"]) or ["Direct"]
            ch_raw = str(ch_list[0] if ch_list else "Direct").strip()
            ch_lower = ch_raw.lower()
            
            # 1. Phân loại Khách hàng (Mục 3)
            is_agency = any(a in ch_lower for a in agency_channels)
            channel_type = "Đại lý" if is_agency else "Khách lẻ"

            code_ev = str(c.get("Code EV", "") or "").strip()
            nat = normalize_nationality(c.get("Quốc tịch (National)"), name, code_ev)

            # 2. Phân loại Loại Visa (Mục 3: Cấu trúc 3 nhóm)
            srv_type = str(c.get("Loại dịch vụ (type)", [""])[0] if c.get("Loại dịch vụ (type)") else "")
            srv_list = [str(s).lower() for s in (c.get("Loại dịch vụ (type)") or [])]
            note_str = str(c.get("Ghi chú", "") or "")
            
            fee_cam_ai = c.get("Lệ phí visa Cam for AI")
            fee_cam_real = c.get("Lệ phí Visa Cam thực tế")
            is_cam = (any("visa cambodia" in s or "visa cam" in s for s in srv_list) and (fee_cam_ai is not None or fee_cam_real is not None)) or ("BUSBY" in name.upper())
            is_multi = ("multi" in note_str.lower()) or any("multi" in s for s in srv_list)
            is_free_vn = ("free visa" in srv_type.lower()) or ("BUSBY" in name.upper())

            if is_cam:
                visa_group = "Visa_Campuchia"
                visa_type_label = "Visa Campuchia"
            elif is_multi:
                visa_group = "Multi"
                visa_type_label = "Multi"
            else:
                visa_group = "Single"
                visa_type_label = "Single"

            # 3. Tính toán Chi phí theo 6 bước (Mục 4)
            # Bước 1: Doanh thu CRM & Tổng tiền HĐ
            try:
                sales_crm = int(str(c.get("Sales revenue", "0") or "0").replace(".", "").replace(",", ""))
            except:
                sales_crm = 0

            if is_agency:
                total_amount = int(round(sales_crm * 1.08)) if sales_crm > 0 else 0
            else:
                total_amount = sales_crm

            # Bước 2: Lệ phí Nhà nước Việt Nam (Mục 2)
            if is_free_vn:
                state_fee_vn = 0
            elif is_multi:
                state_fee_vn = 1325000  # $50
            else:
                state_fee_vn = 662500   # $25

            # Bước 3: Lệ phí Visa Campuchia
            if is_cam:
                if fee_cam_ai is not None:
                    try: fee_cam = int(str(fee_cam_ai).replace(".", "").replace(",", ""))
                    except: fee_cam = 1000000
                elif fee_cam_real is not None:
                    try: fee_cam = int(str(fee_cam_real).replace(".", "").replace(",", ""))
                    except: fee_cam = 1000000
                else:
                    fee_cam = 1000000
            else:
                fee_cam = 0

            # Bước 4: Chi phí Vận tải (Phí xe)
            note_lower = note_str.lower()
            if any(k in srv_list for k in ["90d - bo y", "free visa - bo y", "bờ y", "bo y"]) or "bo y" in note_lower or "bờ y" in note_lower:
                transport_fee = 1250000
            elif any(k in srv_list for k in ["90d - cambodia", "hcm > cam", "free visa - cam", "mộc bài", "moc bai"]) or "mộc bài" in note_lower or "moc bai" in note_lower:
                transport_fee = 1290000
            else:
                transport_fee = 0

            # Bước 5: Phí Dịch vụ gộp (Dòng 1 trong Hợp đồng)
            service_fee_gross = total_amount - state_fee_vn

            # Bước 6: Số tiền hoàn lại khi rớt Visa (Điều 5.2)
            # Hoàn phí Điều 5.2 = Phí dịch vụ gộp (Mục 1) - Phí vận tải - Lệ phí Visa Cam
            refund_amount = max(0, service_fee_gross - transport_fee - fee_cam)

            # Format ngày hạch toán
            acc_ts = c.get("Accounting Date")
            if acc_ts:
                dt = datetime.fromtimestamp(acc_ts / 1000)
                date_vi = dt.strftime("ngày %d tháng %m năm %Y")
                date_en = dt.strftime("%B %d, %Y")
                date_str = dt.strftime("%d/%m/%Y")
            else:
                date_vi = "ngày 16 tháng 08 năm 2026"
                date_en = "August 16, 2026"
                date_str = ""

            safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', name.strip())
            passport_no = extract_passport_no(c, name, code_ev)

            # Tải ảnh hộ chiếu và trích xuất chữ ký
            sig_path = os.path.join("extracted_signatures", f"{safe_name}_{passport_no}_sig.png")
            passports = c.get("Ảnh hộ chiếu")
            if not os.path.exists(sig_path) and passports and isinstance(passports, list) and len(passports) > 0:
                p_item = passports[0]
                ft = p_item.get("file_token")
                p_path = os.path.join("downloads/passports", f"{safe_name}_{passport_no}.jpg")
                if not os.path.exists(p_path) and ft:
                    try:
                        dl_url = f"https://open.larksuite.com/open-apis/drive/v1/medias/{ft}/download"
                        r_dl = await client.get(dl_url, headers=headers)
                        if r_dl.status_code == 200:
                            with open(p_path, "wb") as img_f:
                                img_f.write(r_dl.content)
                    except Exception:
                        pass
                if os.path.exists(p_path):
                    extract_and_clean_signature(p_path, sig_path)

            contract_no = f"{idx:03d}"

            contract_data = {
                "contract_no": contract_no,
                "date_vi": date_vi,
                "date_en": date_en,
                "customer_name": name.upper(),
                "passport_no": passport_no,
                "date_of_issue": "20/05/2022",
                "nationality": nat,
                "total_amount": total_amount,
                "service_fee": refund_amount,
                "transport_fee": transport_fee,
                "state_fee": state_fee_vn,
                "signature_image_path": sig_path if os.path.exists(sig_path) else None
            }

            prefix_type = "Khach_Le" if not is_agency else "Dai_Ly"
            docx_filename = f"{idx:03d}_Hop_Dong_{prefix_type}_{safe_name}.docx"
            pdf_filename = f"{idx:03d}_Hop_Dong_{prefix_type}_{safe_name}.pdf"

            # Xác định thư mục con theo Mục 6
            if not is_agency:
                if is_cam:
                    target_pdf_dir = dir_pdf_kl_cam
                    target_docx_dir = dir_docx_kl_cam
                elif is_multi:
                    target_pdf_dir = dir_pdf_kl_multi
                    target_docx_dir = dir_docx_kl_multi
                else:
                    target_pdf_dir = dir_pdf_kl_single
                    target_docx_dir = dir_docx_kl_single
            else:
                if is_multi:
                    target_pdf_dir = dir_pdf_dl_multi
                    target_docx_dir = dir_docx_dl_multi
                else:
                    target_pdf_dir = dir_pdf_dl_single
                    target_docx_dir = dir_docx_dl_single

            docx_target_path = os.path.join(target_docx_dir, docx_filename)
            pdf_target_path = os.path.join(target_pdf_dir, pdf_filename)

            # 1. Tạo file Word .docx
            build_contract_docx(docx_target_path, contract_data)
            shutil.copyfile(docx_target_path, os.path.join("output_docx_contracts", docx_filename))

            # 2. Tạo file PDF .pdf
            try:
                build_contract_pdf(pdf_target_path, contract_data)
                shutil.copyfile(pdf_target_path, os.path.join("output_contracts", pdf_filename))
            except Exception as e_pdf:
                print(f"⚠️ Lỗi tạo PDF cho {name}: {e_pdf}")

            processed_data.append({
                "stt": f"{idx:03d}",
                "name": name.upper(),
                "passport_no": passport_no,
                "nationality": nat,
                "accounting_date": date_str,
                "channel": ch_raw,
                "channel_type": channel_type,
                "visa_type": visa_type_label,
                "service": srv_type or visa_type_label,
                "sales_crm": sales_crm,
                "total_amount": total_amount,
                "state_fee_vn": state_fee_vn,
                "service_fee_gross": service_fee_gross,
                "transport_fee": transport_fee,
                "fee_cam": fee_cam,
                "refund_amount": refund_amount,
                "pdf_filename": pdf_filename
            })

    print(f"🎉 Đã tạo thành công {len(processed_data)} hợp đồng .docx và .pdf theo đúng cấu trúc thư mục PDF/DOCX!")

    # -------------------------------------------------------------
    # MỤC 6. TẠO FILE EXCEL ĐỐI SOÁT 17 CỘT CHUẨN XÁC 100%
    # -------------------------------------------------------------
    excel_path = "danh_sach_hop_dong_01_08_den_24_08.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active if wb.active is not None else wb.create_sheet()
    ws.title = "Danh_Sach_Hop_Dong"

    # Title
    ws.merge_cells("A1:Q1")
    t_cell = ws["A1"]
    t_cell.value = "BẢNG TỔNG HỢP ĐỐI SOÁT HỢP ĐỒNG KHÁCH LẺ VÀ ĐẠI LÝ (01/08/2026 - 24/08/2026)"
    t_cell.font = Font(name="Arial", size=13, bold=True, color="FFFFFF")
    t_cell.alignment = Alignment(horizontal="center", vertical="center")
    t_cell.fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    ws.row_dimensions[1].height = 35

    ws.merge_cells("A2:Q2")
    s_cell = ws["A2"]
    s_cell.value = f"Tổng số: {len(processed_data)} hồ sơ | Khách lẻ: {sum(1 for x in processed_data if x['channel_type'] == 'Khách lẻ')} | Đại lý: {sum(1 for x in processed_data if x['channel_type'] == 'Đại lý')} | Xuất ngày: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    s_cell.font = Font(name="Arial", size=10, italic=True, color="333333")
    s_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 20

    # 17 CỘT CHUẨN XÁC THEO MỤC 6 CỦA PDF
    headers_17 = [
        "STT",
        "Họ và tên Bên A",
        "Số Hộ Chiếu",
        "Quốc Tịch",
        "Ngày Hạch Toán",
        "Kênh Nguồn",
        "Phân Loại Khách Hàng",
        "Phân Loại Visa",
        "Tuyến / Dịch Vụ",
        "Doanh Thu CRM (VNĐ)",
        "Tổng Tiền HĐ (VNĐ)",
        "Lệ Phí Nhà Nước VN (Mục 2)",
        "Phí Dịch Vụ Gộp (Mục 1)",
        "Chi Phí Vận Tải (Phí xe)",
        "Lệ Phí Visa Cam",
        "Số Tiền Hoàn Lại Khi Rớt Visa (Điều 5.2)",
        "Tên File Hợp Đồng"
    ]

    ws.row_dimensions[3].height = 30
    for col_idx, h in enumerate(headers_17, 1):
        c_hdr = ws.cell(row=3, column=col_idx, value=h)
        c_hdr.font = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
        c_hdr.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c_hdr.fill = PatternFill(start_color="2E75B6", end_color="2E75B6", fill_type="solid")
        c_hdr.border = Border(top=Side(style="thin", color="000000"), bottom=Side(style="thin", color="000000"),
                              left=Side(style="thin", color="000000"), right=Side(style="thin", color="000000"))

    thin_border = Border(top=Side(style="thin", color="D9D9D9"), bottom=Side(style="thin", color="D9D9D9"),
                         left=Side(style="thin", color="D9D9D9"), right=Side(style="thin", color="D9D9D9"))

    for idx, row in enumerate(processed_data, 1):
        r_idx = 3 + idx
        ws.row_dimensions[r_idx].height = 20

        values = [
            row["stt"],
            row["name"],
            row["passport_no"],
            row["nationality"],
            row["accounting_date"],
            row["channel"],
            row["channel_type"],
            row["visa_type"],
            row["service"],
            row["sales_crm"],
            row["total_amount"],
            row["state_fee_vn"],
            row["service_fee_gross"],
            row["transport_fee"],
            row["fee_cam"],
            row["refund_amount"],
            row["pdf_filename"]
        ]

        for col_idx, v in enumerate(values, 1):
            cell = ws.cell(row=r_idx, column=col_idx, value=v)
            cell.font = Font(name="Arial", size=9)
            cell.border = thin_border
            
            if col_idx in [1, 3, 4, 5, 6, 7, 8]:
                cell.alignment = Alignment(horizontal="center", vertical="center")
            elif col_idx in [10, 11, 12, 13, 14, 15, 16]:
                cell.alignment = Alignment(horizontal="right", vertical="center")
                cell.number_format = "#,##0"
            else:
                cell.alignment = Alignment(horizontal="left", vertical="center")

    # Auto-adjust column widths
    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(int(col[0].column or 1))
        ws.column_dimensions[col_letter].width = max(max_len + 3, 12)

    wb.save(excel_path)
    print(f"📊 Đã tạo file Excel đối soát 17 cột: '{excel_path}'")
    return processed_data


if __name__ == "__main__":
    asyncio.run(generate_contracts_by_accounting())

